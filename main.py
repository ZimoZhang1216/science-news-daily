#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import html
import json
import logging
import os
import re
import shutil
import smtplib
import subprocess
import sys
import tempfile
import time
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta, timezone
from email.message import EmailMessage
from email.utils import getaddresses
from pathlib import Path
from typing import Any, Callable
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse

MISSING_DEPENDENCIES: list[str] = []

try:
    import feedparser
except ModuleNotFoundError:
    feedparser = None  # type: ignore[assignment]
    MISSING_DEPENDENCIES.append("feedparser")

try:
    import requests
    from requests.adapters import HTTPAdapter
    from urllib3.util.retry import Retry
except ModuleNotFoundError:
    requests = None  # type: ignore[assignment]
    HTTPAdapter = None  # type: ignore[assignment]
    Retry = None  # type: ignore[assignment]
    MISSING_DEPENDENCIES.append("requests")

try:
    from bs4 import BeautifulSoup
except ModuleNotFoundError:
    BeautifulSoup = None  # type: ignore[assignment]
    MISSING_DEPENDENCIES.append("beautifulsoup4")

try:
    from dateutil import parser as dateparser
except ModuleNotFoundError:
    dateparser = None  # type: ignore[assignment]
    MISSING_DEPENDENCIES.append("python-dateutil")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt, RGBColor
except ModuleNotFoundError:
    Document = None  # type: ignore[assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[assignment]
    Inches = None  # type: ignore[assignment]
    OxmlElement = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    RELATIONSHIP_TYPE = None  # type: ignore[assignment]
    Pt = None  # type: ignore[assignment]
    RGBColor = None  # type: ignore[assignment]
    MISSING_DEPENDENCIES.append("python-docx")

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover - handled at runtime with a clear warning.
    OpenAI = None  # type: ignore[assignment]

try:
    from dotenv import load_dotenv
except ModuleNotFoundError:
    load_dotenv = None  # type: ignore[assignment]
    MISSING_DEPENDENCIES.append("python-dotenv")

if load_dotenv is not None:
    load_dotenv()


LOGGER = logging.getLogger("science_news_daily")

DEFAULT_OUTPUT_DIR = "./output"
DEFAULT_MAX_ITEMS = 30
DEFAULT_MAX_AI_ITEMS = 30
DEFAULT_OPENAI_MODEL = "gpt-5.4-mini"
DEFAULT_DEEPSEEK_MODEL = "deepseek-v4-flash"
DEFAULT_REPORT_EMAIL_TO = "2510248@mail.nankai.edu.cn"
SUPPORTED_LLM_PROVIDERS = {"openai", "deepseek"}
BANNED_TITLE_WORDS = {"震惊", "颠覆", "炸裂", "封神", "逆天", "重磅", "神作", "史诗级"}
USER_AGENT = (
    "ScienceNewsDaily/1.0 "
    "(mailto:please-set-CROSSREF_MAILTO@example.com; Python requests)"
)
DEFAULT_TITLE_STYLE_GUIDE = (
    "中文标题采用严谨的学术亮点风格，优先突出题名或摘要中可核对的研究对象、方法、机制、模型、"
    "数据类型、材料/体系或证据边界。必须以期刊/来源名开头；不要使用公众号悬念、反问、拟人化比喻、"
    "“谁说”“不是……而是……”“一个……让……”等营销句式。"
    "不添加输入里没有的团队、作者、性能数值、临床/产业承诺或因果结论。"
    "标题约 18-36 个中文字符；保留必要英文缩写、化学式、基因名、蛋白名和模型名。"
)
ATTRACTIVE_TITLE_PROMPT_TEMPLATE = (
    "请根据以下论文信息生成一个严谨的中文学术亮点标题。\n"
    "要求：\n"
    "1. 必须保留期刊/来源名作为标题开头。\n"
    "2. 标题主体突出研究对象、方法、机制、模型、数据类型、材料/体系或证据边界。\n"
    "3. 标题长度控制在 18-36 个中文字符。\n"
    "4. 不使用反问、悬念、夸张、拟人化比喻和营销号句式。\n"
    "5. 不使用“谁说”“不是……而是……”“一个……让……”等表达。\n"
    "6. 不夸大结论，不制造虚假因果，不加入输入中没有的数值或应用承诺。\n\n"
    "论文信息：\n"
    "来源：{source}\n"
    "英文标题：{title}\n"
    "摘要：{abstract}\n"
    "领域：{field}\n\n"
    "输出只给一个标题。"
)

CHEMISTRY_TERM_TRANSLATIONS: list[tuple[str, str]] = [
    ("click chemistry-enabled", "点击化学"),
    ("click chemistry", "点击化学"),
    ("food safety", "食品安全"),
    ("low-barrier hydrogen bond", "低势垒氢键"),
    ("hydrogen bond", "氢键"),
    ("radical transfer", "自由基转移"),
    ("electron transfer", "电子转移"),
    ("proton supply", "质子供给"),
    ("decoupling", "解耦"),
    ("oxidized-state accumulation", "氧化态累积"),
    ("water oxidation kinetics", "水氧化动力学"),
    ("water oxidation", "水氧化"),
    ("active site", "活性位"),
    ("photoredox", "光氧化还原"),
    ("visible-light", "可见光"),
    ("photoinduced", "光诱导"),
    ("photochemical", "光化学"),
    ("electrocatal", "电催化"),
    ("photocatal", "光催化"),
    ("catalyst", "催化剂"),
    ("catalysis", "催化"),
    ("weakly solvating", "弱溶剂"),
    ("electrolyte", "电解液"),
    ("battery", "电池"),
    ("batteries", "电池"),
    ("solar cell", "太阳能电池"),
    ("total synthesis", "全合成"),
    ("asymmetric synthesis", "不对称合成"),
    ("cross-coupling", "交叉偶联"),
    ("cyclization", "环化"),
    ("metal-organic framework", "MOF"),
    ("metal organic framework", "MOF"),
    ("covalent organic framework", "COF"),
    ("porous", "多孔材料"),
    ("framework", "框架"),
    ("polymer", "聚合物"),
    ("plastic", "塑料"),
    ("recycling", "回收"),
    ("fluorescence", "荧光"),
    ("phosphorescence", "磷光"),
    ("luminescence", "发光"),
    ("spectroscopy", "光谱"),
    ("sensor", "传感"),
    ("analytical", "分析检测"),
    ("density functional", "DFT"),
    ("machine learning", "机器学习"),
    ("computational", "计算模拟"),
    ("simulation", "模拟"),
    ("autoencoder", "自编码器"),
    ("randomforest", "随机森林"),
    ("metabolomics", "代谢组学"),
    ("protein", "蛋白"),
    ("enzyme", "酶"),
]

BIOLOGY_TERM_TRANSLATIONS: list[tuple[str, str]] = [
    ("esophageal squamous cell carcinoma", "食管鳞状细胞癌"),
    ("intrahepatic cholangiocarcinoma", "肝内胆管癌"),
    ("cholangiocarcinoma", "胆管癌"),
    ("super-enhancer-associated", "超级增强子"),
    ("super-enhancer-driven", "超级增强子"),
    ("super-enhancer", "超级增强子"),
    ("long non-coding rna", "lncRNA"),
    ("lncrna", "lncRNA"),
    ("mir205hg", "MIR205HG"),
    ("glycolysis reprogramming", "糖酵解重编程"),
    ("glycolysis", "糖酵解"),
    ("pyroptosis", "细胞焦亡"),
    ("synthetic lethality", "合成致死"),
    ("rna-triggered", "RNA触发"),
    ("cell killing", "细胞杀伤"),
    ("cas12a2", "Cas12a2"),
    ("dna shredding", "DNA广泛切割"),
    ("gene expression", "基因表达"),
    ("mutation-bearing", "突变细胞"),
    ("low-barrier hydrogen bond", "低势垒氢键"),
    ("hydrogen bond", "氢键"),
    ("radical transfer", "自由基转移"),
    ("single-cell transcriptome", "单细胞转录组"),
    ("single-cell transcriptomic", "单细胞转录组"),
    ("spatial transcriptomic", "空间转录组"),
    ("spatial transcriptome", "空间转录组"),
    ("bulk transcriptome", "bulk转录组"),
    ("transcriptome", "转录组"),
    ("transcriptional regulation", "转录调控"),
    ("transcriptional", "转录调控"),
    ("phosphoproteome", "磷酸化蛋白组"),
    ("proteome", "蛋白组"),
    ("tumor microenvironment", "肿瘤微环境"),
    ("tumor progression", "肿瘤进展"),
    ("cancer progression", "癌症进展"),
    ("therapeutic applications", "治疗应用"),
    ("immune checkpoint", "免疫检查点"),
    ("t cell", "T细胞"),
    ("macrophage", "巨噬细胞"),
    ("immune", "免疫"),
    ("immunity", "免疫"),
    ("neuron", "神经元"),
    ("synapse", "突触"),
    ("brain", "大脑"),
    ("astrocyte", "星形胶质细胞"),
    ("microbiome", "微生物组"),
    ("bacteria", "细菌"),
    ("microbial", "微生物"),
    ("viral", "病毒"),
    ("virus", "病毒"),
    ("genome", "基因组"),
    ("crispr", "CRISPR"),
    ("epigen", "表观遗传"),
    ("protein", "蛋白"),
    ("enzyme", "酶"),
    ("receptor", "受体"),
    ("ligand", "配体"),
    ("stem cell", "干细胞"),
    ("organoid", "类器官"),
    ("development", "发育"),
    ("embryo", "胚胎"),
    ("metabolism", "代谢"),
    ("mitochondria", "线粒体"),
    ("clinical", "临床"),
    ("patient", "患者"),
    ("cancer", "癌症"),
    ("tumor", "肿瘤"),
    ("disease", "疾病机制"),
    ("therapy", "治疗线索"),
]

BIOLOGY_METHOD_TRANSLATIONS: list[tuple[str, str]] = [
    ("single-cell transcriptome", "单细胞转录组"),
    ("single-cell transcriptomic", "单细胞转录组"),
    ("spatial transcriptomic", "空间转录组"),
    ("spatial transcriptome", "空间转录组"),
    ("bulk transcriptome", "bulk转录组"),
    ("transcriptome", "转录组"),
    ("phosphoproteome", "磷酸化蛋白组"),
    ("proteome", "蛋白组"),
    ("genomic mutation", "基因组突变数据"),
    ("chip-seq", "ChIP-Seq"),
    ("hichip-seq", "HiChIP-Seq"),
    ("chip-qpcr", "ChIP-qPCR"),
    ("dual-luciferase", "双荧光素酶报告实验"),
    ("rna-seq", "RNA-seq"),
    ("scrna-seq", "scRNA-seq"),
    ("reporter assay", "报告基因实验"),
    ("qpcr", "qPCR"),
    ("western blot", "Western blot"),
    ("crispr", "CRISPR筛选"),
]

BIOLOGY_BROAD_TERMS = {
    "疾病机制",
    "治疗线索",
    "临床",
    "患者",
    "发育",
    "蛋白",
    "基因组",
}

BIOLOGY_TITLE_EVIDENCE: list[tuple[str, tuple[str, ...]]] = [
    ("CRISPR", ("crispr", "cas12", "cas9", "cas13")),
    ("Cas12a2", ("cas12a2",)),
    ("病毒", ("virus", "viral", "hiv", "sars", "infect")),
    ("表观遗传", ("epigen", "chromatin", "histone", "methylation")),
    ("临床", ("clinical", "patient", "trial", "cohort")),
    ("转录组", ("transcriptome", "transcriptomic", "rna-seq", "scrna")),
    ("单细胞", ("single-cell", "single cell", "scrna")),
    ("空间转录组", ("spatial transcript",)),
    ("T细胞", ("t cell", "t-cell")),
    ("肠道菌", ("microbiome", "gut", "bacteria", "microbial")),
]

STATISTICS_TERM_TRANSLATIONS: list[tuple[str, str]] = [
    ("bayesian", "贝叶斯"),
    ("causal inference", "因果推断"),
    ("high-dimensional", "高维数据"),
    ("time series", "时间序列"),
    ("spatial", "空间统计"),
    ("survival", "生存分析"),
    ("random forest", "随机森林"),
    ("deep learning", "深度学习"),
    ("machine learning", "机器学习"),
    ("expectation-maximization", "EM算法"),
    ("em algorithm", "EM算法"),
    ("bootstrap", "Bootstrap"),
    ("prediction", "预测"),
    ("parameter", "参数"),
    ("confidence interval", "置信区间"),
    ("regression", "回归"),
    ("missing data", "缺失数据"),
    ("meta-analysis", "Meta分析"),
]

PROFILE_TERM_TRANSLATIONS = {
    "chemistry": CHEMISTRY_TERM_TRANSLATIONS,
    "biology": BIOLOGY_TERM_TRANSLATIONS,
    "statistics": STATISTICS_TERM_TRANSLATIONS,
}

GENERIC_TITLE_BODIES = {
    "不是催化越复杂，而是活性位更关键",
    "一个弱溶剂策略，让快充电池更耐温",
    "从碎片到骨架，合成路线有了新搭法",
    "一个多孔框架，让分子筛选更精准",
    "不是塑料难回收，而是化学键要会断",
    "发光的关键，可能藏在分子振动里",
    "不是只靠试错，而是模型先探路",
    "一个检测平台，让微量信号看得见",
    "把化学工具送进细胞，信号读得更准",
    "不是T细胞太弱，而是微环境还在踩刹车",
    "一条皮层通路，让声音行为有了专线",
    "肠道菌的全球流行病，可能藏在基因组里",
    "一个基因开关，让调控图谱再更新",
    "蛋白机器的关键，可能藏在隐藏档位",
    "从单个细胞到类器官，命运轨迹更清楚",
    "一条代谢暗线，把生命过程串起来",
    "疾病机制的关键，可能藏在新按钮里",
}

FIELD_KEYWORDS: dict[str, list[str]] = {
    "有机化学": [
        "organic synthesis",
        "total synthesis",
        "asymmetric synthesis",
        "photoredox",
        "cross-coupling",
        "C-H activation",
        "organocatalysis",
        "stereoselective",
        "natural product",
        "synthetic methodology",
    ],
    "物理化学": [
        "physical chemistry",
        "spectroscopy",
        "kinetics",
        "thermodynamics",
        "photochemistry",
        "electrochemistry",
        "ultrafast",
        "surface chemistry",
        "quantum chemistry",
        "molecular dynamics",
    ],
    "材料化学": [
        "materials chemistry",
        "polymer",
        "perovskite",
        "metal-organic framework",
        "MOF",
        "COF",
        "nanomaterial",
        "2D material",
        "semiconductor",
        "self-assembly",
    ],
    "化学生物学": [
        "chemical biology",
        "bioorthogonal",
        "proteomics",
        "drug discovery",
        "enzyme",
        "protein",
        "peptide",
        "biosynthesis",
        "metabolite",
        "bioconjugation",
    ],
    "催化": [
        "catalysis",
        "catalyst",
        "electrocatalysis",
        "photocatalysis",
        "heterogeneous catalysis",
        "homogeneous catalysis",
        "single-atom catalyst",
        "organometallic",
        "activation",
    ],
    "能源化学": [
        "energy chemistry",
        "battery",
        "lithium",
        "sodium-ion",
        "hydrogen",
        "fuel cell",
        "CO2 reduction",
        "oxygen evolution",
        "solar cell",
        "water splitting",
    ],
    "计算化学": [
        "computational chemistry",
        "density functional theory",
        "DFT",
        "molecular simulation",
        "machine learning",
        "quantum computation",
        "molecular docking",
        "ab initio",
        "force field",
    ],
    "分析化学": [
        "analytical chemistry",
        "mass spectrometry",
        "chromatography",
        "sensor",
        "biosensor",
        "imaging",
        "NMR",
        "Raman",
        "fluorescence",
        "single-cell",
    ],
}

CHEMISTRY_TERMS = sorted(
    {
        keyword.lower()
        for keywords in FIELD_KEYWORDS.values()
        for keyword in keywords
    }
    | {
        "chemistry",
        "chemical",
        "molecule",
        "molecular",
        "reaction",
        "synthesis",
        "catalytic",
        "polymerization",
        "electrolyte",
        "redox",
    }
)

LEARNING_VALUE_TERMS: dict[str, float] = {
    "review": 8.0,
    "perspective": 7.0,
    "viewpoint": 5.0,
    "tutorial": 6.0,
    "mechanism": 5.0,
    "mechanistic": 5.0,
    "design principle": 5.0,
    "benchmark": 4.5,
    "platform": 4.0,
    "general method": 4.0,
    "scalable": 3.5,
    "selective": 3.0,
    "high-throughput": 3.0,
    "structure-property": 3.0,
    "insight": 2.5,
    "framework": 2.5,
    "strategy": 2.0,
    "methodology": 2.0,
}

ARXIV_QUERY_TERMS = [
    "chemistry",
    "catalysis",
    "electrochemistry",
    "organic synthesis",
    "materials chemistry",
    "chemical biology",
    "computational chemistry",
    "physical chemistry",
    "analytical chemistry",
    "battery",
    "polymer",
    "perovskite",
]

PUBMED_QUERY_TERMS = [
    "chemical biology",
    "medicinal chemistry",
    "organic synthesis",
    "catalysis",
    "analytical chemistry",
    "mass spectrometry",
    "metabolomics",
    "drug discovery",
    "bioorthogonal chemistry",
    "chemical proteomics",
]

CROSSREF_JOURNALS: list[dict[str, Any]] = [
    {"source": "JACS", "issns": ["0002-7863", "1520-5126"], "broad": False},
    {
        "source": "Angewandte Chemie",
        "issns": ["1433-7851", "1521-3773"],
        "broad": False,
    },
    {"source": "Nature Chemistry", "issns": ["1755-4330", "1755-4349"], "broad": False},
    {"source": "Science", "issns": ["0036-8075", "1095-9203"], "broad": True},
    {"source": "ACS Catalysis", "issns": ["2155-5435"], "broad": False},
    {"source": "ACS Energy Letters", "issns": ["2380-8195"], "broad": False},
    {"source": "ACS Materials Letters", "issns": ["2639-4979"], "broad": False},
    {"source": "ACS Central Science", "issns": ["2374-7943"], "broad": False},
    {"source": "Organic Letters", "issns": ["1523-7060", "1523-7052"], "broad": False},
    {
        "source": "The Journal of Organic Chemistry",
        "issns": ["0022-3263", "1520-6904"],
        "broad": False,
    },
    {"source": "Analytical Chemistry", "issns": ["0003-2700", "1520-6882"], "broad": False},
    {
        "source": "The Journal of Physical Chemistry Letters",
        "issns": ["1948-7185"],
        "broad": False,
    },
    {"source": "Chemical Science", "issns": ["2041-6520", "2041-6539"], "broad": False},
    {
        "source": "Chemical Society Reviews",
        "issns": ["0306-0012", "1460-4744"],
        "broad": False,
    },
    {
        "source": "Journal of Materials Chemistry A",
        "issns": ["2050-7488", "2050-7496"],
        "broad": False,
    },
    {
        "source": "Energy & Environmental Science",
        "issns": ["1754-5692", "1754-5706"],
        "broad": False,
    },
    {
        "source": "Catalysis Science & Technology",
        "issns": ["2044-4753", "2044-4761"],
        "broad": False,
    },
    {
        "source": "Physical Chemistry Chemical Physics",
        "issns": ["1463-9076", "1463-9084"],
        "broad": False,
    },
]

RSS_FEEDS: list[dict[str, Any]] = [
    {
        "source": "C&EN (ACS)",
        "url": "https://feeds.feedburner.com/cen_latestnews",
        "broad": False,
    },
    {
        "source": "Nature Chemistry",
        "url": "https://www.nature.com/nchem.rss",
        "broad": False,
    },
    {
        "source": "Science",
        "url": "https://www.science.org/action/showFeed?type=etoc&feed=rss&jc=science",
        "broad": True,
    },
    {
        "source": "Chemistry World News (RSC)",
        "url": "https://www.chemistryworld.com/409.rss",
        "broad": False,
    },
    {
        "source": "Chemistry World Research (RSC)",
        "url": "https://www.chemistryworld.com/410.rss",
        "broad": False,
    },
]

SOURCE_WEIGHTS = {
    "Nature Chemistry": 80,
    "JACS": 78,
    "Angewandte Chemie": 76,
    "Science": 74,
    "Chemical Science": 68,
    "Chemical Society Reviews": 68,
    "ACS Central Science": 66,
    "ACS Catalysis": 64,
    "ACS Energy Letters": 64,
    "Energy & Environmental Science": 64,
    "PubMed": 55,
    "arXiv": 48,
}

BIOLOGY_FIELD_KEYWORDS: dict[str, list[str]] = {
    "分子生物学": [
        "molecular biology",
        "protein",
        "RNA",
        "DNA",
        "transcription",
        "translation",
        "chromatin",
        "enzyme",
        "signaling",
        "structural biology",
    ],
    "细胞生物学": [
        "cell",
        "organelle",
        "membrane",
        "stem cell",
        "cell cycle",
        "single-cell",
        "development",
        "apoptosis",
        "differentiation",
    ],
    "遗传与基因组": [
        "genomics",
        "genome",
        "genetics",
        "variant",
        "GWAS",
        "epigenomics",
        "sequencing",
        "CRISPR",
        "gene editing",
    ],
    "免疫学": [
        "immunology",
        "immune",
        "T cell",
        "B cell",
        "antibody",
        "vaccine",
        "inflammation",
        "autoimmune",
        "tumor immunity",
    ],
    "神经科学": [
        "neuroscience",
        "neuron",
        "synapse",
        "brain",
        "neural circuit",
        "neurodegeneration",
        "cognition",
        "glia",
    ],
    "生物技术": [
        "biotechnology",
        "synthetic biology",
        "protein engineering",
        "cell therapy",
        "gene therapy",
        "screening",
        "organoid",
        "spatial transcriptomics",
    ],
    "微生物与进化": [
        "microbiome",
        "microbial",
        "bacteria",
        "virus",
        "pathogen",
        "evolution",
        "ecology",
        "host-microbe",
    ],
    "疾病机制": [
        "cancer",
        "disease",
        "metabolism",
        "pathogenesis",
        "therapeutic",
        "drug target",
        "biomarker",
        "clinical",
    ],
}

BIOLOGY_TERMS = sorted(
    {
        keyword.lower()
        for keywords in BIOLOGY_FIELD_KEYWORDS.values()
        for keyword in keywords
    }
    | {
        "biology",
        "biological",
        "gene",
        "protein",
        "cellular",
        "organism",
        "tissue",
        "molecular",
        "life science",
    }
)

BIOLOGY_ARXIV_QUERY_TERMS = [
    "cat:q-bio.BM",
    "cat:q-bio.CB",
    "cat:q-bio.GN",
    "cat:q-bio.MN",
    "cat:q-bio.NC",
    "cat:q-bio.PE",
    "cat:q-bio.QM",
    "cat:q-bio.SC",
    "cat:q-bio.TO",
    "single-cell",
    "genomics",
    "synthetic biology",
    "protein design",
]

BIOLOGY_PUBMED_QUERY_TERMS = [
    "molecular biology",
    "cell biology",
    "genomics",
    "single-cell",
    "CRISPR",
    "immunology",
    "neuroscience",
    "synthetic biology",
    "protein engineering",
    "microbiome",
    "cancer biology",
    "gene therapy",
]

BIOLOGY_CROSSREF_JOURNALS: list[dict[str, Any]] = [
    {"source": "Nature", "issns": ["0028-0836", "1476-4687"], "broad": True},
    {"source": "Science", "issns": ["0036-8075", "1095-9203"], "broad": True},
    {"source": "Cell", "issns": ["0092-8674", "1097-4172"], "broad": False},
    {"source": "Nature Biotechnology", "issns": ["1087-0156", "1546-1696"], "broad": False},
    {"source": "Nature Methods", "issns": ["1548-7091", "1548-7105"], "broad": False},
    {"source": "Nature Genetics", "issns": ["1061-4036", "1546-1718"], "broad": False},
    {"source": "Nature Medicine", "issns": ["1078-8956", "1546-170X"], "broad": True},
    {"source": "Molecular Cell", "issns": ["1097-2765", "1097-4164"], "broad": False},
    {"source": "Genome Biology", "issns": ["1474-760X"], "broad": False},
    {"source": "PLOS Biology", "issns": ["1545-7885", "1544-9173"], "broad": False},
    {"source": "eLife", "issns": ["2050-084X"], "broad": False},
    {"source": "PNAS", "issns": ["0027-8424", "1091-6490"], "broad": True},
]

BIOLOGY_RSS_FEEDS: list[dict[str, Any]] = [
    {"source": "Nature", "url": "https://www.nature.com/nature.rss", "broad": True},
    {"source": "Nature Biotechnology", "url": "https://www.nature.com/nbt.rss", "broad": False},
    {"source": "Nature Methods", "url": "https://www.nature.com/nmeth.rss", "broad": False},
    {"source": "Nature Genetics", "url": "https://www.nature.com/ng.rss", "broad": False},
    {"source": "Nature Medicine", "url": "https://www.nature.com/nm.rss", "broad": True},
    {
        "source": "Science",
        "url": "https://www.science.org/action/showFeed?type=etoc&feed=rss&jc=science",
        "broad": True,
    },
    {"source": "Cell", "url": "https://www.cell.com/cell/current.rss", "broad": False},
]

BIOLOGY_SOURCE_WEIGHTS = {
    "Nature Biotechnology": 72,
    "Nature Methods": 70,
    "Nature Genetics": 70,
    "Nature Medicine": 70,
    "Nature": 78,
    "Science": 76,
    "Molecular Cell": 66,
    "Cell": 76,
    "Genome Biology": 64,
    "PLOS Biology": 62,
    "eLife": 60,
    "PNAS": 58,
    "PubMed": 55,
    "arXiv": 48,
}

STATISTICS_FIELD_KEYWORDS: dict[str, list[str]] = {
    "统计理论": [
        "asymptotic",
        "estimator",
        "estimation",
        "inference",
        "hypothesis testing",
        "confidence interval",
        "minimax",
        "efficiency",
    ],
    "贝叶斯统计": [
        "bayesian",
        "posterior",
        "prior",
        "MCMC",
        "variational inference",
        "hierarchical model",
        "probabilistic",
    ],
    "因果推断": [
        "causal inference",
        "treatment effect",
        "instrumental variable",
        "counterfactual",
        "difference-in-differences",
        "mediation",
    ],
    "高维统计": [
        "high-dimensional",
        "sparse",
        "regularization",
        "lasso",
        "graphical model",
        "multiple testing",
        "dimension reduction",
    ],
    "机器学习统计": [
        "statistical learning",
        "machine learning",
        "deep learning",
        "generalization",
        "uncertainty quantification",
        "prediction",
        "classification",
    ],
    "时间序列与空间统计": [
        "time series",
        "spatial statistics",
        "spatio-temporal",
        "forecasting",
        "stochastic process",
        "state space",
    ],
    "生物统计": [
        "biostatistics",
        "clinical trial",
        "survival analysis",
        "epidemiology",
        "competing risks",
        "longitudinal data",
        "statistical genetics",
    ],
    "计算统计": [
        "computational statistics",
        "simulation",
        "bootstrap",
        "Monte Carlo",
        "sampling",
        "optimization",
        "algorithm",
    ],
}

STATISTICS_TERMS = sorted(
    {
        keyword.lower()
        for keywords in STATISTICS_FIELD_KEYWORDS.values()
        for keyword in keywords
    }
    | {
        "statistics",
        "statistical",
        "probability",
        "model",
        "data analysis",
        "regression",
        "distribution",
        "uncertainty",
    }
)

STATISTICS_ARXIV_QUERY_TERMS = [
    "cat:stat.AP",
    "cat:stat.CO",
    "cat:stat.ME",
    "cat:stat.ML",
    "cat:stat.OT",
    "cat:stat.TH",
    "cat:math.ST",
    "causal inference",
    "bayesian statistics",
    "high-dimensional statistics",
    "time series",
]

STATISTICS_PUBMED_QUERY_TERMS = [
    "biostatistics",
    "survival analysis",
    "clinical trial statistics",
    "causal inference",
    "statistical genetics",
    "longitudinal data analysis",
]

STATISTICS_CROSSREF_JOURNALS: list[dict[str, Any]] = [
    {"source": "Annals of Statistics", "issns": ["0090-5364", "2168-8966"], "broad": False},
    {"source": "Biometrika", "issns": ["0006-3444", "1464-3510"], "broad": False},
    {"source": "Journal of the American Statistical Association", "issns": ["0162-1459", "1537-274X"], "broad": False},
    {"source": "JRSS Series B", "issns": ["1369-7412", "1467-9868"], "broad": False},
    {"source": "Statistical Science", "issns": ["0883-4237", "2168-8745"], "broad": False},
    {"source": "Bayesian Analysis", "issns": ["1936-0975"], "broad": False},
    {"source": "Bernoulli", "issns": ["1350-7265", "1573-9759"], "broad": False},
    {"source": "Annals of Applied Statistics", "issns": ["1932-6157", "1941-7330"], "broad": False},
    {"source": "Journal of Machine Learning Research", "issns": ["1532-4435"], "broad": False},
]

STATISTICS_RSS_FEEDS: list[dict[str, Any]] = [
    {
        "source": "arXiv stat",
        "url": "https://rss.arxiv.org/rss/stat",
        "broad": False,
    },
]

STATISTICS_SOURCE_WEIGHTS = {
    "Annals of Statistics": 76,
    "Biometrika": 74,
    "Journal of the American Statistical Association": 74,
    "JRSS Series B": 72,
    "Statistical Science": 68,
    "Bayesian Analysis": 66,
    "Bernoulli": 64,
    "Annals of Applied Statistics": 62,
    "Journal of Machine Learning Research": 60,
    "PubMed": 50,
    "arXiv": 52,
}

REPORT_PROFILES: dict[str, dict[str, Any]] = {
    "chemistry": {
        "key": "chemistry",
        "title": "化学科研资讯日报",
        "failure_title": "化学科研资讯日报运行失败报告",
        "output_prefix": "chem_news",
        "header_label": "CHEM NEWS DAILY",
        "meta_fields": "有机、物化、材料、化学生物、催化、能源、计算、分析",
        "field_keywords": FIELD_KEYWORDS,
        "relevance_terms": CHEMISTRY_TERMS,
        "arxiv_query_terms": ARXIV_QUERY_TERMS,
        "pubmed_query_terms": PUBMED_QUERY_TERMS,
        "crossref_journals": CROSSREF_JOURNALS,
        "rss_feeds": RSS_FEEDS,
        "source_weights": SOURCE_WEIGHTS,
        "default_field": "综合化学",
        "ai_role": "化学领域科研编辑",
        "ai_task": "生成化学科研资讯日报摘要",
        "title_style": (
            "化学标题取消营销号和公众号悬念风格，严谨性优先。标题应突出可从题名或摘要确认的学术亮点，"
            "包括研究对象、反应/材料/催化体系、方法学创新、机制问题、表征或计算框架。"
            "优先使用“期刊名：研究对象/方法/机制的学术亮点”结构，例如"
            "“JACS：铱催化羰基参与的不对称氢芳基化”、"
            "“Angew：氧化态累积调控水氧化动力学”、"
            "“Nature Chemistry：DNA条形码适配体库用于化学多样性筛选”。"
            "不要使用反问句、悬念句、拟人化比喻或“不是……而是……”“一个……让……”等营销句式；"
            "不添加输入里没有的团队、作者、性能数值、应用承诺或因果结论。"
        ),
        "email_env": "CHEM_REPORT_EMAIL_TO",
        "default_email_to": DEFAULT_REPORT_EMAIL_TO,
    },
    "biology": {
        "key": "biology",
        "title": "生物科研资讯日报",
        "failure_title": "生物科研资讯日报运行失败报告",
        "output_prefix": "bio_news",
        "header_label": "BIO NEWS DAILY",
        "meta_fields": "分子、细胞、基因组、免疫、神经、生物技术、微生物、疾病机制",
        "field_keywords": BIOLOGY_FIELD_KEYWORDS,
        "relevance_terms": BIOLOGY_TERMS,
        "arxiv_query_terms": BIOLOGY_ARXIV_QUERY_TERMS,
        "pubmed_query_terms": BIOLOGY_PUBMED_QUERY_TERMS,
        "crossref_journals": BIOLOGY_CROSSREF_JOURNALS,
        "rss_feeds": BIOLOGY_RSS_FEEDS,
        "source_weights": BIOLOGY_SOURCE_WEIGHTS,
        "default_field": "综合生物学",
        "ai_role": "生物学领域科研编辑",
        "ai_task": "生成生物科研资讯日报摘要",
        "title_style": (
            "生物标题与化学日报对齐，采用严谨的学术亮点风格。标题应突出题名或摘要中可核对的"
            "细胞类型、基因/蛋白/RNA、疾病或表型、实验方法、数据类型、模型系统或机制对象。"
            "优先使用“期刊名：研究对象/方法/机制的学术亮点”结构，例如"
            "“Science：短RNA伴侣调控TDP-43聚集构象”、"
            "“Nature：CRISPR-Cas12a介导RNA触发的细胞杀伤”、"
            "“Genome Biology：单细胞图谱解析免疫微环境异质性”。"
            "不要使用反问句、悬念句、拟人化比喻或“不是……而是……”“一个……让……”等营销句式；"
            "不添加输入里没有的治疗效果、人群结论、临床承诺或因果外推。"
        ),
        "email_env": "BIO_REPORT_EMAIL_TO",
        "default_email_to": "",
    },
    "statistics": {
        "key": "statistics",
        "title": "统计学科研资讯日报",
        "failure_title": "统计学科研资讯日报运行失败报告",
        "output_prefix": "stat_news",
        "header_label": "STAT NEWS DAILY",
        "meta_fields": "理论、贝叶斯、因果、高维、机器学习、时空、生统、计算统计",
        "field_keywords": STATISTICS_FIELD_KEYWORDS,
        "relevance_terms": STATISTICS_TERMS,
        "arxiv_query_terms": STATISTICS_ARXIV_QUERY_TERMS,
        "pubmed_query_terms": STATISTICS_PUBMED_QUERY_TERMS,
        "crossref_journals": STATISTICS_CROSSREF_JOURNALS,
        "rss_feeds": STATISTICS_RSS_FEEDS,
        "source_weights": STATISTICS_SOURCE_WEIGHTS,
        "default_field": "综合统计学",
        "ai_role": "统计学领域科研编辑",
        "ai_task": "生成统计学科研资讯日报摘要",
        "title_style": (
            "统计学标题与化学日报对齐，采用严谨的学术亮点风格。标题应突出题名或摘要中可核对的"
            "统计问题、模型类别、估计/检验方法、理论性质、算法框架、数据结构或应用边界。"
            "优先使用“期刊名/来源：问题/方法/理论性质的学术亮点”结构，例如"
            "“Ann. Stat.：高维模型中的稳健推断界限”、"
            "“JASA：缺失数据下的因果效应估计”、"
            "“arXiv：扩散模型评估的非参数检验框架”。"
            "不要使用反问句、悬念句、拟人化比喻或“不是……而是……”“一个……让……”等营销句式；"
            "不添加输入里没有的性能数值、应用承诺或理论结论。"
        ),
        "email_env": "STAT_REPORT_EMAIL_TO",
        "default_email_to": "",
    },
}


@dataclass
class NewsItem:
    title: str
    source: str
    published: datetime | None
    link: str
    abstract: str = ""
    doi: str = ""
    authors: list[str] = field(default_factory=list)
    field_name: str = "综合化学"
    item_id: str = ""
    attractive_title: str = ""
    chinese_title: str = ""
    comment: str = ""
    score: float = 0.0


@dataclass
class SourceStatus:
    name: str
    success: bool
    item_count: int = 0
    error: str = ""

    def summary(self) -> str:
        if self.success:
            return f"成功，获取 {self.item_count} 条"
        return f"失败：{summarize_source_error(self.error)}"


@dataclass
class LLMConfig:
    provider: str
    model: str
    api_key: str
    api_key_env: str
    base_url: str | None = None


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    text = html.unescape(str(value))
    if "<" in text and ">" in text and BeautifulSoup is not None:
        text = BeautifulSoup(text, "html.parser").get_text(" ", strip=True)
    text = (
        text.replace("\u00a0", " ")
        .replace("\u202f", " ")
        .replace("\u2007", " ")
        .replace("\u2009", " ")
        .replace("\u2010", "-")
        .replace("\u2011", "-")
        .replace("\u2012", "-")
        .replace("\u2212", "-")
        .replace("\ufeff", "")
    )
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def summarize_source_error(error: str) -> str:
    text = clean_text(error)
    if not text:
        return "未知错误。"
    lowered = text.lower()
    if "read timed out" in lowered or "readtimeouterror" in lowered or "timeout" in lowered:
        return "请求超时，来源响应过慢；其他来源已继续抓取，可稍后重试。"
    if "name or service not known" in lowered or "failed to resolve" in lowered or "nameresolutionerror" in lowered:
        return "DNS 解析失败；请检查网络、代理或 DNS 配置。"
    if "connection refused" in lowered or "connectionerror" in lowered:
        return "连接失败；可能是来源临时不可用或本地网络波动。"
    if "429" in lowered or "too many requests" in lowered:
        return "触发限流；建议降低请求频率或稍后重试。"
    if "403" in lowered or "forbidden" in lowered:
        return "访问被拒绝；该来源可能限制了自动化请求。"
    if "500" in lowered or "502" in lowered or "503" in lowered or "504" in lowered:
        return "来源服务器返回 5xx 错误；建议稍后重试。"
    return truncate(text, 180)


def parse_datetime(value: Any) -> datetime | None:
    if not value:
        return None
    if dateparser is None:
        return None
    try:
        parsed = dateparser.parse(str(value))
    except (TypeError, ValueError, OverflowError):
        return None
    if not parsed:
        return None
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def date_parts_to_datetime(value: dict[str, Any] | None) -> datetime | None:
    if not value:
        return None
    parts_list = value.get("date-parts") or []
    if not parts_list or not parts_list[0]:
        return None
    parts = parts_list[0]
    try:
        year = int(parts[0])
        month = int(parts[1]) if len(parts) > 1 else 1
        day = int(parts[2]) if len(parts) > 2 else 1
        return datetime(year, month, day, tzinfo=timezone.utc)
    except (TypeError, ValueError):
        return None


def format_date(value: datetime | None) -> str:
    if not value:
        return "未知日期"
    return value.astimezone().strftime("%Y-%m-%d")


def truncate(text: str, limit: int) -> str:
    text = clean_text(text)
    if len(text) <= limit:
        return text
    return text[: limit - 1].rstrip() + "..."


def stable_index(value: str, modulo: int, offset: int = 0) -> int:
    if modulo <= 0:
        return 0
    digest = hashlib.sha1(clean_text(value).encode("utf-8")).hexdigest()
    return (int(digest[:8], 16) + offset) % modulo


def normalize_chinese_title(title: str) -> str:
    normalized = clean_text(title)
    normalized = re.sub(r"\s*[:：]\s*", "：", normalized)
    normalized = re.sub(r"\s*[—–]\s*", "——", normalized)
    normalized = re.sub(r"\s+-\s+", "——", normalized)
    normalized = re.sub(r"—{3,}", "——", normalized)
    normalized = re.sub(r"\s+([，。；：！？、）】》])", r"\1", normalized)
    normalized = re.sub(r"([，。；：！？、])\s+", r"\1", normalized)
    normalized = re.sub(r"([（【《])\s+", r"\1", normalized)
    return normalized


def chinese_char_count(text: str) -> int:
    return sum(1 for char in text if "\u4e00" <= char <= "\u9fff")


def normalize_doi(value: str) -> str:
    doi = clean_text(value)
    doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", doi, flags=re.IGNORECASE)
    doi = re.sub(r"^doi\s*[:：]\s*", "", doi, flags=re.IGNORECASE)
    doi = doi.strip().strip(".;,)")
    return doi.lower()


def extract_doi(value: str) -> str:
    text = clean_text(value)
    match = re.search(r"10\.\d{4,9}/[-._;()/:A-Z0-9]+", text, flags=re.IGNORECASE)
    if not match:
        return ""
    return normalize_doi(match.group(0))


def canonical_url_key(value: str) -> str:
    url = clean_text(value)
    if not url:
        return ""
    parsed = urlparse(url)
    scheme = (parsed.scheme or "https").lower()
    netloc = parsed.netloc.lower()
    if netloc.startswith("www."):
        netloc = netloc[4:]
    path = re.sub(r"/+$", "", parsed.path or "/")
    if netloc in {"doi.org", "dx.doi.org"}:
        doi = normalize_doi(path.lstrip("/"))
        return f"doi:{doi}" if doi else ""
    ignored_query_keys = {
        "utm_source",
        "utm_medium",
        "utm_campaign",
        "utm_term",
        "utm_content",
        "fbclid",
        "gclid",
    }
    query_pairs = [
        (key, val)
        for key, val in parse_qsl(parsed.query, keep_blank_values=True)
        if key.lower() not in ignored_query_keys
    ]
    query = urlencode(sorted(query_pairs), doseq=True)
    return urlunparse((scheme, netloc, path, "", query, ""))


def title_fingerprint(value: str) -> str:
    text = clean_text(value).lower()
    text = re.sub(r"^(article|research article|review|abstract)\s*[:：-]\s*", "", text)
    text = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", text)
    return text


def item_identity_keys(item: NewsItem) -> list[str]:
    keys: list[str] = []
    doi = normalize_doi(item.doi) or extract_doi(item.link)
    if doi:
        keys.append(f"doi:{doi}")
    url_key = canonical_url_key(item.link)
    if url_key:
        keys.append(url_key if url_key.startswith("doi:") else f"url:{url_key}")
    title_key = title_fingerprint(item.title)
    if len(title_key) >= 18:
        keys.append(f"title:{title_key}")
    if not keys:
        keys.append(f"fallback:{stable_index(item.title + item.source + item.link, 10**8)}")
    return list(dict.fromkeys(keys))


def duplicate_quality_score(item: NewsItem) -> float:
    abstract = clean_text(item.abstract)
    score = 0.0
    if normalize_doi(item.doi) or extract_doi(item.link):
        score += 45
    if item.link:
        score += 10
    if item.published:
        score += 6
    if item.authors:
        score += min(len(item.authors), 5)
    if abstract and "未提供摘要" not in abstract and "RSS 未提供摘要" not in abstract:
        score += min(len(abstract) / 8, 55)
    else:
        score -= 12
    source = clean_text(item.source).lower()
    if "crossref" not in source and not source.startswith("rss"):
        score += 3
    return score


def infer_profile_key(item: NewsItem, profile: dict[str, Any] | None = None) -> str:
    if profile and profile.get("key"):
        return str(profile["key"])
    field_name = item.field_name
    if any(token in field_name for token in ("统计", "贝叶斯", "因果", "高维", "生统")):
        return "statistics"
    if any(token in field_name for token in ("生物", "分子", "细胞", "免疫", "神经", "基因")):
        return "biology"
    return "chemistry"


def field_short_name(field_name: str) -> str:
    field_name = clean_text(field_name)
    if not field_name:
        return "研究对象"
    replacements = {
        "综合化学": "化学问题",
        "综合生物学": "生命机制",
        "综合统计学": "统计问题",
        "化学生物学": "化学生物",
        "材料化学": "材料体系",
        "物理化学": "物化机制",
        "能源化学": "能源材料",
        "计算化学": "计算模型",
        "分析化学": "分析检测",
        "有机化学": "有机反应",
    }
    return replacements.get(field_name, field_name.replace("综合", "") or "研究对象")


def add_unique_term(terms: list[str], term: str, limit: int) -> None:
    term = clean_text(term)
    if not term or len(term) > 18:
        return
    if any(term == existing or (len(term) > 2 and term in existing) or (len(existing) > 2 and existing in term) for existing in terms):
        return
    terms.append(term)
    if len(terms) > limit:
        del terms[limit:]


def extract_cn_terms(item: NewsItem, profile: dict[str, Any] | None = None, limit: int = 3) -> list[str]:
    profile_key = infer_profile_key(item, profile)
    haystack = f"{item.title} {item.abstract}".lower()
    terms: list[str] = []
    translations = PROFILE_TERM_TRANSLATIONS.get(profile_key, CHEMISTRY_TERM_TRANSLATIONS)
    for needle, cn_term in translations:
        if needle.lower() in haystack:
            add_unique_term(terms, cn_term, limit)
        if len(terms) >= limit:
            return terms

    title_tokens = re.findall(r"\b[A-Z][A-Za-z0-9-]{2,}\b|\b[A-Z]{2,}\b", clean_text(item.title))
    token_blacklist = {
        "The",
        "This",
        "With",
        "From",
        "Using",
        "Based",
        "Study",
        "Review",
        "Article",
        "Journal",
        "Science",
        "Nature",
        "Chemistry",
    }
    for token in title_tokens:
        if token in token_blacklist or len(token) > 12:
            continue
        if profile_key == "biology" and not (token.isupper() or re.search(r"\d", token)):
            continue
        add_unique_term(terms, token, limit)
        if len(terms) >= limit:
            return terms

    add_unique_term(terms, field_short_name(item.field_name), limit)
    return terms


def title_body(title: str) -> str:
    normalized = normalize_chinese_title(title)
    if "：" in normalized:
        return normalized.split("：", 1)[1]
    return normalized


def title_body_key(title: str) -> str:
    body = title_body(title)
    body = re.sub(r"\s+", "", body)
    body = re.sub(r"[，。；：！？、,.!?;:]", "", body)
    return body


def is_generic_title(title: str) -> bool:
    body = title_body(title)
    return body in GENERIC_TITLE_BODIES or title_body_key(title) in {
        re.sub(r"[，。；：！？、,.!?;:]", "", body) for body in GENERIC_TITLE_BODIES
    }


def comment_key(comment: str) -> str:
    text = clean_text(comment)
    text = re.sub(r"“[^”]{8,}”", "“X”", text)
    text = re.sub(r"[A-Za-z][A-Za-z0-9 ,;:()/_\\.-]{18,}", "X", text)
    text = re.sub(r"\d+(?:\.\d+)?", "0", text)
    text = re.sub(r"\s+", "", text)
    return text[:120]


def is_mostly_english_text(text: str) -> bool:
    normalized = clean_text(text)
    ascii_chars = sum(1 for char in normalized if ord(char) < 128 and char.isalpha())
    letter_chars = sum(1 for char in normalized if char.isalpha())
    return bool(letter_chars and ascii_chars / letter_chars > 0.72)


def is_low_value_comment(comment: str) -> bool:
    text = clean_text(comment)
    low_value_markers = (
        "英文摘要显示其围绕",
        "建议结合原文核对材料体系",
        "建议结合原文核对",
        "主要内容为：出版商元数据未提供摘要",
        "主要内容为：RSS 未提供摘要",
        "题名和摘要能确认本条围绕",
        "可靠元数据指向",
        "可核对线索集中在",
        "从题名和摘要看，工作关注",
    )
    return any(marker in text for marker in low_value_markers)


def is_marketing_title(title: str) -> bool:
    body = title_body(title)
    marketing_markers = (
        "谁说",
        "有戏",
        "故事",
        "解锁",
        "露出",
        "藏在",
        "给答案",
        "更完整",
        "更清楚",
        "换个看法",
        "遇上",
        "交给",
        "替身",
        "搭起",
        "让",
        "不是",
    )
    if "？" in body or "?" in body:
        return True
    return any(marker in body for marker in marketing_markers)


def biology_title_has_unsupported_terms(title: str, item: NewsItem) -> bool:
    body = title_body(title)
    haystack = f"{item.title} {item.abstract}".lower()
    for cn_term, evidence_terms in BIOLOGY_TITLE_EVIDENCE:
        if cn_term in body and not any(evidence in haystack for evidence in evidence_terms):
            return True
    return False


def extract_terms_from_pairs(text: str, pairs: list[tuple[str, str]], limit: int) -> list[str]:
    haystack = clean_text(text).lower()
    terms: list[str] = []
    for needle, cn_term in pairs:
        if needle.lower() in haystack:
            add_unique_term(terms, cn_term, limit)
        if len(terms) >= limit:
            break
    return terms


def extract_biology_focus_terms(item: NewsItem, limit: int = 5) -> list[str]:
    terms = extract_cn_terms(item, REPORT_PROFILES["biology"], limit=limit + 4)
    specific = [term for term in terms if term not in BIOLOGY_BROAD_TERMS]
    if specific:
        terms = specific

    title_tokens = re.findall(r"\b[A-Z][A-Za-z0-9-]{3,}\b|\b[A-Z]{2,}\b", clean_text(item.title))
    for token in title_tokens:
        if token in {"RNA", "DNA", "METHODS", "RESULTS", "BACKGROUND"}:
            continue
        if re.search(r"\d", token) or token.isupper():
            add_unique_term(terms, token, limit + 2)
    return terms[:limit]


def extract_biology_method_terms(item: NewsItem, limit: int = 4) -> list[str]:
    return extract_terms_from_pairs(item.abstract, BIOLOGY_METHOD_TRANSLATIONS, limit)


def format_cn_terms(terms: list[str], fallback: str) -> str:
    if not terms:
        return fallback
    if len(terms) == 1:
        return terms[0]
    return "、".join(terms)


def fallback_biology_comment(item: NewsItem, variant_offset: int = 0) -> str:
    abstract = clean_text(item.abstract)
    abstract = re.sub(r"^(abstract|summary)\s*[:：]?\s*", "", abstract, flags=re.IGNORECASE)
    abstract = re.sub(
        r"\b(background|objective|purpose|methods?|results?|conclusions?)\s*[:：]\s*",
        "",
        abstract,
        flags=re.IGNORECASE,
    )
    focus_terms = extract_biology_focus_terms(item, limit=5)
    focus_text = format_cn_terms(focus_terms, field_short_name(item.field_name))
    method_terms = extract_biology_method_terms(item, limit=4)
    method_text = format_cn_terms(method_terms, "")

    if abstract and "未提供摘要" not in abstract and "RSS 未提供摘要" not in abstract:
        method_clause = (
            f"摘要提到{method_text}等证据来源"
            if method_text
            else "摘要提供了研究背景和实验线索"
        )
        templates = [
            "研究聚焦{focus}；{method_clause}。日报只概括题名和摘要中可确认的信息，具体机制和适用边界以原文为准。",
            "该工作围绕{focus}展开；{method_clause}，后续阅读可重点核对模型、对照实验和结论外推范围。",
            "题名和摘要指向{focus}相关问题；{method_clause}。这里不扩展未在摘要中出现的临床或应用结论。",
        ]
        index = stable_index(item.title + item.source, len(templates), variant_offset)
        return templates[index].format(focus=focus_text, method_clause=method_clause)

    return (
        f"出版商元数据未给出充分摘要；题名可确认本条关注{focus_text}。"
        "建议打开 DOI 或原始链接核对实验体系、模型和结论边界。"
    )


def build_session() -> requests.Session:
    if requests is None or Retry is None or HTTPAdapter is None:
        raise RuntimeError("requests is not installed")
    session = requests.Session()
    retry = Retry(
        total=2,
        connect=2,
        read=2,
        status=2,
        backoff_factor=0.8,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET"],
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    session.headers.update({"User-Agent": os.getenv("CHEM_NEWS_USER_AGENT", USER_AGENT)})
    return session


def resolve_profile(profile_key: str) -> dict[str, Any]:
    key = (profile_key or "chemistry").strip().lower()
    profile = REPORT_PROFILES.get(key)
    if profile is None:
        supported = ", ".join(sorted(REPORT_PROFILES))
        raise SystemExit(f"--profile must be one of: {supported}")
    return profile


def classify_field(title: str, abstract: str, profile: dict[str, Any]) -> str:
    haystack = f"{title} {abstract}".lower()
    scores: dict[str, int] = {}
    for field_name, keywords in profile["field_keywords"].items():
        score = 0
        for keyword in keywords:
            if keyword.lower() in haystack:
                score += 2 if " " in keyword else 1
        if score:
            scores[field_name] = score
    if not scores:
        return profile["default_field"]
    return max(scores.items(), key=lambda item: item[1])[0]


def is_profile_relevant(item: NewsItem, profile: dict[str, Any]) -> bool:
    haystack = f"{item.title} {item.abstract}".lower()
    return any(term in haystack for term in profile["relevance_terms"])


def text_from_xml(element: ET.Element | None) -> str:
    if element is None:
        return ""
    return clean_text("".join(element.itertext()))


def arxiv_query_fragment(term: str) -> str:
    term = term.strip()
    if term.startswith("cat:"):
        return term
    if " " in term:
        return f'all:"{term}"'
    return f"all:{term}"


def fetch_arxiv(
    session: requests.Session,
    since: datetime,
    until: datetime,
    max_items: int,
    profile: dict[str, Any],
) -> list[NewsItem]:
    search_terms = []
    for term in profile["arxiv_query_terms"]:
        search_terms.append(arxiv_query_fragment(term))
    params = {
        "search_query": " OR ".join(search_terms),
        "start": 0,
        "max_results": min(max(max_items, 10), 100),
        "sortBy": "submittedDate",
        "sortOrder": "descending",
    }
    response = session.get("https://export.arxiv.org/api/query", params=params, timeout=30)
    response.raise_for_status()
    feed = feedparser.parse(response.text)
    items: list[NewsItem] = []
    for entry in feed.entries:
        published = parse_datetime(getattr(entry, "published", None)) or parse_datetime(
            getattr(entry, "updated", None)
        )
        if published and not (since <= published <= until):
            continue
        title = clean_text(getattr(entry, "title", ""))
        abstract = clean_text(getattr(entry, "summary", ""))
        link = clean_text(getattr(entry, "link", ""))
        authors = [clean_text(author.get("name", "")) for author in getattr(entry, "authors", [])]
        if not title or not link:
            continue
        item = NewsItem(
            title=title,
            source="arXiv",
            published=published,
            link=link,
            abstract=abstract,
            authors=[author for author in authors if author],
        )
        item.field_name = classify_field(item.title, item.abstract, profile)
        if is_profile_relevant(item, profile):
            items.append(item)
    return items


def pubmed_params(extra: dict[str, Any]) -> dict[str, Any]:
    params = {"tool": "science_news_daily", **extra}
    if os.getenv("NCBI_EMAIL"):
        params["email"] = os.environ["NCBI_EMAIL"]
    if os.getenv("NCBI_API_KEY"):
        params["api_key"] = os.environ["NCBI_API_KEY"]
    return params


def parse_pubmed_date(article: ET.Element) -> datetime | None:
    article_date = article.find(".//ArticleDate")
    if article_date is not None:
        year = text_from_xml(article_date.find("Year"))
        month = text_from_xml(article_date.find("Month"))
        day = text_from_xml(article_date.find("Day"))
        parsed = parse_datetime("-".join(part for part in [year, month, day] if part))
        if parsed:
            return parsed

    pub_date = article.find(".//Journal/JournalIssue/PubDate")
    if pub_date is None:
        return None
    year = text_from_xml(pub_date.find("Year"))
    month = text_from_xml(pub_date.find("Month"))
    day = text_from_xml(pub_date.find("Day"))
    medline_date = text_from_xml(pub_date.find("MedlineDate"))
    raw = " ".join(part for part in [year, month, day] if part) or medline_date
    return parse_datetime(raw)


def fetch_pubmed(
    session: requests.Session,
    since: datetime,
    until: datetime,
    max_items: int,
    profile: dict[str, Any],
) -> list[NewsItem]:
    query_terms = profile["pubmed_query_terms"]
    if not query_terms:
        return []
    from_date = since.strftime("%Y/%m/%d")
    to_date = until.strftime("%Y/%m/%d")
    term_query = " OR ".join(f'"{term}"[Title/Abstract]' for term in query_terms)
    query = f"({term_query}) AND ({from_date}[PDAT] : {to_date}[PDAT])"
    search_response = session.get(
        "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi",
        params=pubmed_params(
            {
                "db": "pubmed",
                "term": query,
                "retmode": "json",
                "retmax": min(max(max_items, 10), 100),
                "sort": "pub date",
            }
        ),
        timeout=30,
    )
    search_response.raise_for_status()
    id_list = search_response.json().get("esearchresult", {}).get("idlist", [])
    if not id_list:
        return []

    fetch_response = session.get(
        "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi",
        params=pubmed_params({"db": "pubmed", "id": ",".join(id_list), "retmode": "xml"}),
        timeout=45,
    )
    fetch_response.raise_for_status()
    root = ET.fromstring(fetch_response.content)
    items: list[NewsItem] = []
    for article in root.findall(".//PubmedArticle"):
        pmid = text_from_xml(article.find(".//PMID"))
        title = text_from_xml(article.find(".//ArticleTitle"))
        journal = text_from_xml(article.find(".//Journal/Title")) or "PubMed"
        published = parse_pubmed_date(article)
        if published and not (since <= published <= until):
            continue
        abstract_parts = []
        for abstract_element in article.findall(".//Abstract/AbstractText"):
            label = clean_text(abstract_element.attrib.get("Label", ""))
            text = text_from_xml(abstract_element)
            if not text:
                continue
            abstract_parts.append(f"{label}: {text}" if label else text)
        authors = []
        for author in article.findall(".//Author"):
            last_name = text_from_xml(author.find("LastName"))
            initials = text_from_xml(author.find("Initials"))
            full = " ".join(part for part in [last_name, initials] if part)
            if full:
                authors.append(full)
        doi = ""
        for article_id in article.findall(".//ArticleId"):
            if article_id.attrib.get("IdType") == "doi":
                doi = text_from_xml(article_id)
                break
        if not title or not pmid:
            continue
        item = NewsItem(
            title=title,
            source=f"PubMed: {journal}",
            published=published,
            link=f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
            abstract=" ".join(abstract_parts),
            doi=doi,
            authors=authors,
        )
        item.field_name = classify_field(item.title, item.abstract, profile)
        if is_profile_relevant(item, profile):
            items.append(item)
    return items


def crossref_polite_params(extra: dict[str, Any]) -> dict[str, Any]:
    params = dict(extra)
    if os.getenv("CROSSREF_MAILTO"):
        params["mailto"] = os.environ["CROSSREF_MAILTO"]
    return params


def crossref_publication_date(work: dict[str, Any]) -> datetime | None:
    for key in ["published-online", "published-print", "published", "issued", "created"]:
        published = date_parts_to_datetime(work.get(key))
        if published:
            return published
    return None


def fetch_crossref_journal(
    session: requests.Session,
    journal: dict[str, Any],
    since: datetime,
    until: datetime,
    rows_per_issn: int,
    profile: dict[str, Any],
) -> list[NewsItem]:
    items: list[NewsItem] = []
    seen: set[str] = set()
    for issn in journal["issns"]:
        params = crossref_polite_params(
            {
                "filter": (
                    f"from-pub-date:{since.date().isoformat()},"
                    f"until-pub-date:{until.date().isoformat()},"
                    f"type:journal-article,issn:{issn}"
                ),
                "sort": "published",
                "order": "desc",
                "rows": min(max(rows_per_issn, 5), 50),
            }
        )
        response = session.get("https://api.crossref.org/works", params=params, timeout=30)
        response.raise_for_status()
        works = response.json().get("message", {}).get("items", [])
        for work in works:
            doi = clean_text(work.get("DOI", ""))
            if doi and doi.lower() in seen:
                continue
            title = clean_text((work.get("title") or [""])[0])
            if not title:
                continue
            container_title = clean_text((work.get("container-title") or [journal["source"]])[0])
            published = crossref_publication_date(work)
            if published and not (since <= published <= until):
                continue
            abstract = clean_text(work.get("abstract", ""))
            link = clean_text(work.get("URL", "")) or (f"https://doi.org/{doi}" if doi else "")
            author_names = []
            for author in work.get("author", [])[:8]:
                given = clean_text(author.get("given", ""))
                family = clean_text(author.get("family", ""))
                full = " ".join(part for part in [given, family] if part)
                if full:
                    author_names.append(full)
            item = NewsItem(
                title=title,
                source=journal["source"] if journal["source"] in container_title else container_title,
                published=published,
                link=link,
                abstract=abstract or "出版商元数据未提供摘要；请通过链接查看原文摘要。",
                doi=doi,
                authors=author_names,
            )
            item.field_name = classify_field(item.title, item.abstract, profile)
            if not journal.get("broad") or is_profile_relevant(item, profile):
                items.append(item)
                if doi:
                    seen.add(doi.lower())
        time.sleep(0.15)
    return items


def fetch_crossref(
    session: requests.Session, since: datetime, until: datetime, max_items: int, profile: dict[str, Any]
) -> tuple[list[NewsItem], list[SourceStatus]]:
    items: list[NewsItem] = []
    statuses: list[SourceStatus] = []
    journals = profile["crossref_journals"]
    rows_per_issn = max(5, min(20, max_items // max(len(journals), 1) + 2))
    for journal in journals:
        try:
            journal_items = fetch_crossref_journal(session, journal, since, until, rows_per_issn, profile)
            LOGGER.info("%s via Crossref: %d items", journal["source"], len(journal_items))
            items.extend(journal_items)
            statuses.append(
                SourceStatus(
                    name=f"Crossref: {journal['source']}",
                    success=True,
                    item_count=len(journal_items),
                )
            )
        except Exception as exc:  # noqa: BLE001 - one journal must not stop the report.
            LOGGER.warning("Crossref source failed for %s: %s", journal["source"], exc)
            statuses.append(
                SourceStatus(
                    name=f"Crossref: {journal['source']}",
                    success=False,
                    error=f"{type(exc).__name__}: {exc}",
                )
            )
    return items, statuses


def fetch_rss(
    session: requests.Session, since: datetime, until: datetime, max_items: int, profile: dict[str, Any]
) -> tuple[list[NewsItem], list[SourceStatus]]:
    items: list[NewsItem] = []
    statuses: list[SourceStatus] = []
    feeds = profile["rss_feeds"]
    per_feed = max(10, min(40, max_items // max(len(feeds), 1) + 5))
    for feed_config in feeds:
        try:
            response = session.get(feed_config["url"], timeout=30)
            response.raise_for_status()
            parsed = feedparser.parse(response.content)
            count = 0
            for entry in parsed.entries[:per_feed]:
                published = parse_datetime(
                    getattr(entry, "published", None) or getattr(entry, "updated", None)
                )
                if published and not (since <= published <= until):
                    continue
                title = clean_text(getattr(entry, "title", ""))
                abstract = clean_text(
                    getattr(entry, "summary", "") or getattr(entry, "description", "")
                )
                link = clean_text(getattr(entry, "link", ""))
                if not title or not link:
                    continue
                item = NewsItem(
                    title=title,
                    source=feed_config["source"],
                    published=published,
                    link=link,
                    abstract=abstract or "RSS 未提供摘要；请通过链接查看详情。",
                )
                item.field_name = classify_field(item.title, item.abstract, profile)
                if not feed_config.get("broad") or is_profile_relevant(item, profile):
                    items.append(item)
                    count += 1
            LOGGER.info("%s RSS: %d items", feed_config["source"], count)
            statuses.append(
                SourceStatus(
                    name=f"RSS: {feed_config['source']}",
                    success=True,
                    item_count=count,
                )
            )
        except Exception as exc:  # noqa: BLE001 - RSS endpoints are best-effort.
            LOGGER.warning("RSS source failed for %s: %s", feed_config["source"], exc)
            statuses.append(
                SourceStatus(
                    name=f"RSS: {feed_config['source']}",
                    success=False,
                    error=f"{type(exc).__name__}: {exc}",
                )
            )
    return items, statuses


def dedupe_items(items: list[NewsItem]) -> list[NewsItem]:
    clusters: list[list[NewsItem]] = []
    key_to_cluster: dict[str, int] = {}
    for item in items:
        keys = item_identity_keys(item)
        matched_clusters = sorted({key_to_cluster[key] for key in keys if key in key_to_cluster})
        if not matched_clusters:
            cluster_index = len(clusters)
            clusters.append([item])
            for key in keys:
                key_to_cluster[key] = cluster_index
            continue

        target_index = matched_clusters[0]
        clusters[target_index].append(item)
        for duplicate_index in reversed(matched_clusters[1:]):
            clusters[target_index].extend(clusters[duplicate_index])
            clusters[duplicate_index] = []
            for key, cluster_index in list(key_to_cluster.items()):
                if cluster_index == duplicate_index:
                    key_to_cluster[key] = target_index
        for key in keys:
            key_to_cluster[key] = target_index

    unique_items: list[NewsItem] = []
    for cluster in clusters:
        if not cluster:
            continue
        unique_items.append(max(cluster, key=duplicate_quality_score))
    return unique_items


def rank_item(item: NewsItem, now: datetime, profile: dict[str, Any]) -> float:
    source_weight = 45
    for source_prefix, weight in profile["source_weights"].items():
        if item.source.startswith(source_prefix) or source_prefix in item.source:
            source_weight = weight
            break
    haystack = f"{item.title} {item.abstract}".lower()
    keyword_hits = sum(1 for term in profile["relevance_terms"] if term in haystack)
    learning_bonus = sum(weight for term, weight in LEARNING_VALUE_TERMS.items() if term in haystack)
    abstract_bonus = min(len(item.abstract) / 450, 6)
    title_bonus = 4 if any(term in item.title.lower() for term in LEARNING_VALUE_TERMS) else 0
    metadata_penalty = 5 if "未提供摘要" in item.abstract or not item.abstract else 0
    recency_bonus = 0.0
    if item.published:
        age_hours = max((now - item.published).total_seconds() / 3600, 0)
        recency_bonus = max(0, 72 - age_hours) / 72 * 10
    return (
        source_weight
        + keyword_hits * 1.4
        + learning_bonus
        + title_bonus
        + abstract_bonus
        + recency_bonus
        - metadata_penalty
    )


def prepare_items(items: list[NewsItem], max_items: int, now: datetime, profile: dict[str, Any]) -> list[NewsItem]:
    unique = dedupe_items(items)
    for item in unique:
        item.field_name = classify_field(item.title, item.abstract, profile)
        item.score = rank_item(item, now, profile)
    ranked = sorted(
        unique,
        key=lambda item: (
            item.score,
            item.published.timestamp() if item.published else 0,
            item.title,
        ),
        reverse=True,
    )
    for index, item in enumerate(ranked[:max_items], start=1):
        item.item_id = f"N{index:03d}"
    return ranked[:max_items]


def fallback_comment(item: NewsItem, variant_offset: int = 0) -> str:
    if infer_profile_key(item) == "biology":
        return fallback_biology_comment(item, variant_offset)

    abstract = clean_text(item.abstract)
    abstract = re.sub(r"^(abstract|summary)\s*[:：]?\s*", "", abstract, flags=re.IGNORECASE)
    terms = extract_cn_terms(item, limit=3)
    term_text = "、".join(terms) if terms else f"{item.field_name}相关问题"
    if abstract and "未提供摘要" not in abstract:
        ascii_chars = sum(1 for char in abstract if ord(char) < 128 and char.isalpha())
        letter_chars = sum(1 for char in abstract if char.isalpha())
        if letter_chars and ascii_chars / letter_chars > 0.72:
            templates = [
                "题名和摘要能确认本条围绕{terms}展开；为避免过度解读，具体体系、指标和结论边界以原文为准。",
                "可靠元数据指向{terms}相关研究；摘要信息有限时，日报只保留可确认对象，不延伸未给出的应用结论。",
                "从题名和摘要看，工作关注{terms}；后续阅读可重点核对实验条件、模型假设和关键评价指标。",
                "本条的可核对线索集中在{terms}；判断其重要性时，建议优先查看原文图表、对照和限制说明。",
            ]
            index = stable_index(item.title + item.source, len(templates), variant_offset)
            return templates[index].format(terms=term_text)
        abstract = truncate(abstract, 120)
        return f"该条目聚焦{item.field_name}方向，摘要显示其主要内容为：{abstract}"
    templates = [
        "出版商元数据较少，但题名可确认其与{terms}有关；建议打开 DOI 或原始链接核对摘要和全文。",
        "本条只保留来源、日期和题名中的可靠线索：{terms}；详细方法和结论边界需以原文为准。",
        "摘要暂缺，日报不扩展未给出的发现；可先从{terms}入手判断是否值得阅读全文。",
    ]
    index = stable_index(item.title + item.source, len(templates), variant_offset)
    return templates[index].format(terms=term_text)


def normalize_comment(comment: str, item: NewsItem) -> str:
    normalized = clean_text(comment)
    if not normalized:
        return fallback_comment(item)
    without_prefix = re.sub(r"^(abstract|summary)\s*[:：]?\s*", "", normalized, flags=re.IGNORECASE)
    without_prefix = re.sub(r"^[（(]\s*与\s*N\d+\s*相同\s*[）)]\s*", "", without_prefix)
    without_prefix = re.sub(r"^[（(]\s*信息有限\s*[）)]\s*", "信息有限：", without_prefix)
    if "ABSTRACT" in normalized.upper() or is_mostly_english_text(without_prefix):
        return fallback_comment(item)
    return without_prefix


def source_title_prefix(source: str, profile: dict[str, Any] | None = None) -> str:
    normalized = clean_text(source)
    if not normalized:
        return "科研快讯"
    pubmed_match = re.match(r"^PubMed\s*[:：]\s*(.+)$", normalized, flags=re.IGNORECASE)
    pubmed_journal = clean_text(pubmed_match.group(1)) if pubmed_match else ""
    if profile and profile.get("key") == "biology":
        lowered = (pubmed_journal or normalized).lower()
        if lowered == "nature":
            return "Nature"
        if lowered.startswith("nature "):
            return "Nature子刊"
        biology_prefixes = {
            "Cell": "Cell子刊",
            "Science": "Science",
            "Proceedings of the National Academy of Sciences": "PNAS",
            "PNAS": "PNAS",
            "The Lancet": "Lancet",
            "New England Journal of Medicine": "NEJM",
            "Genome Biology": "Genome Biology",
            "eLife": "eLife",
        }
        for needle, prefix in biology_prefixes.items():
            if needle.lower() in lowered:
                return prefix
    known_prefixes = {
        "Journal of the American Chemical Society": "JACS",
        "JACS": "JACS",
        "Angewandte Chemie International Edition": "Angew",
        "Angewandte Chemie": "Angew",
        "Proceedings of the National Academy of Sciences": "PNAS",
        "PNAS": "PNAS",
        "Nature Chemistry": "Nature Chemistry",
        "Nature": "Nature",
        "Science": "Science",
        "Chemical Science": "Chem. Sci.",
        "ACS Catalysis": "ACS Catalysis",
        "ACS Central Science": "ACS Central Science",
        "Critical Reviews in Analytical Chemistry": "Crit. Rev. Anal. Chem.",
        "Analytical Chemistry": "Anal. Chem.",
        "The Journal of Organic Chemistry": "JOC",
        "Organic Letters": "Org. Lett.",
        "The Journal of Physical Chemistry Letters": "JPCL",
    }
    matching_text = f"{pubmed_journal} {normalized}".strip()
    for needle, prefix in known_prefixes.items():
        if needle.lower() in matching_text.lower():
            return prefix
    if pubmed_journal:
        if len(pubmed_journal) <= 24 and "journal" not in pubmed_journal.lower():
            return truncate(pubmed_journal, 22)
        return "PubMed"
    return truncate(re.sub(r"\s+via\s+.*$", "", normalized, flags=re.IGNORECASE), 22)


def title_terms_pair(item: NewsItem, profile: dict[str, Any]) -> tuple[str, str]:
    terms = extract_cn_terms(item, profile, limit=3)
    field_term = field_short_name(item.field_name)
    first = terms[0] if terms else field_term
    second = ""
    for term in terms[1:] + [field_term, "原文证据"]:
        if term and term != first and term not in first and first not in term:
            second = term
            break
    return first, second or "原文证据"


def rule_based_chinese_title(item: NewsItem, profile: dict[str, Any], variant_offset: int = 0) -> str:
    prefix = source_title_prefix(item.source, profile)
    first, second = title_terms_pair(item, profile)
    seed = f"{item.title} {item.source}"
    profile_key = profile.get("key")

    if second == "原文证据":
        if profile_key == "biology":
            templates_one_term = [
                "{a}相关机制的近期进展",
                "{a}研究中的方法学线索",
                "{a}方向的新近论文",
                "{a}过程的证据边界",
            ]
        elif profile_key == "statistics":
            templates_one_term = [
                "{a}问题的统计方法进展",
                "{a}模型的理论线索",
                "{a}方向的新近论文",
                "{a}分析的推断框架",
            ]
        else:
            templates_one_term = [
                "{a}相关体系的近期进展",
                "{a}研究中的方法学线索",
                "{a}方向的新近论文",
                "{a}问题的机制研究",
            ]
        template = templates_one_term[stable_index(seed, len(templates_one_term), variant_offset)]
        return normalize_chinese_title(f"{prefix}：{template.format(a=first)}")

    if profile_key == "biology":
        templates = [
            "{a}相关{b}机制研究",
            "{a}模型中的{b}证据",
            "{a}方法用于{b}解析",
            "{a}与{b}调控关系",
            "{a}介导的{b}过程",
            "{a}方向的{b}研究进展",
            "{a}过程中的{b}表征线索",
        ]
    elif profile_key == "statistics":
        templates = [
            "{a}模型中的{b}推断",
            "{a}方法用于{b}分析",
            "{a}问题的{b}框架",
            "{a}与{b}统计建模",
            "{a}方向的{b}理论线索",
            "{a}算法的{b}评估",
            "{a}数据中的{b}分析",
        ]
    else:
        templates = [
            "{a}相关{b}机制研究",
            "{a}体系中的{b}调控",
            "{a}方法用于{b}分析",
            "{a}与{b}耦合机制",
            "{a}介导的{b}过程",
            "{a}方向的{b}研究进展",
            "{a}体系的{b}表征线索",
        ]
    template = templates[stable_index(seed, len(templates), variant_offset)]
    return normalize_chinese_title(f"{prefix}：{template.format(a=first, b=second)}")


def fallback_chinese_title(item: NewsItem, profile: dict[str, Any]) -> str:
    existing_title = normalize_chinese_title(item.attractive_title or item.chinese_title)
    if (
        existing_title
        and not any(word in existing_title for word in BANNED_TITLE_WORDS)
        and not is_generic_title(existing_title)
        and not is_marketing_title(existing_title)
        and not (profile.get("key") == "biology" and biology_title_has_unsupported_terms(existing_title, item))
        and chinese_char_count(existing_title) <= 46
    ):
        return existing_title
    return rule_based_chinese_title(item, profile)


def compact_title_source_prefix(title: str, source_alias: str) -> str:
    parts = [part.strip() for part in normalize_chinese_title(title).split("：") if part.strip()]
    if len(parts) <= 2:
        return "：".join(parts)

    def is_source_like(part: str) -> bool:
        lowered = part.lower()
        letters = sum(1 for char in part if char.isalpha())
        non_space = sum(1 for char in part if not char.isspace())
        mostly_latin = bool(non_space) and letters / non_space > 0.55
        return (
            source_alias.lower() in lowered
            or lowered.startswith("pubmed")
            or lowered in {"nature", "science", "cell", "jacs", "angew"}
            or ("journal" in lowered and len(part) <= 34)
            or (mostly_latin and chinese_char_count(part) == 0 and len(part) <= 28)
        )

    body_index = 1
    while body_index < len(parts) - 1 and is_source_like(parts[body_index]):
        body_index += 1
    if body_index > 1:
        return normalize_chinese_title(f"{source_alias}：{'：'.join(parts[body_index:])}")
    return "：".join(parts)


def normalize_attractive_title(title: str, item: NewsItem, profile: dict[str, Any]) -> str:
    normalized = normalize_chinese_title(title)
    fallback = fallback_chinese_title(item, profile)
    source_alias = source_title_prefix(item.source, profile)
    if not normalized:
        return fallback
    if any(word in normalized for word in BANNED_TITLE_WORDS):
        return fallback

    source_part = normalized.split("：", 1)[0]
    if source_alias not in source_part and source_alias not in normalized[: max(len(source_alias) + 3, 12)]:
        normalized = normalize_chinese_title(f"{source_alias}：{normalized}")

    normalized = compact_title_source_prefix(normalized, source_alias)

    if is_generic_title(normalized):
        return rule_based_chinese_title(item, profile)
    if is_marketing_title(normalized):
        return rule_based_chinese_title(item, profile)
    if profile.get("key") == "biology" and biology_title_has_unsupported_terms(normalized, item):
        return rule_based_chinese_title(item, profile)
    if chinese_char_count(normalized) > 46:
        return fallback
    return normalized


def display_title(item: NewsItem) -> str:
    return item.attractive_title or item.chinese_title or item.title


def ensure_unique_titles_and_comments(items: list[NewsItem], profile: dict[str, Any]) -> None:
    seen_titles: set[str] = set()
    seen_comments: set[str] = set()
    for item in items:
        title = normalize_attractive_title(item.attractive_title or item.chinese_title, item, profile)
        for offset in range(12):
            key = title_body_key(title)
            if key and key not in seen_titles and not is_generic_title(title):
                break
            title = rule_based_chinese_title(item, profile, offset + 1)
        item.attractive_title = title
        item.chinese_title = item.chinese_title or title
        seen_titles.add(title_body_key(title))

        comment = normalize_comment(item.comment, item)
        for offset in range(12):
            key = comment_key(comment)
            if key and key not in seen_comments and not is_low_value_comment(comment):
                break
            comment = fallback_comment(item, offset + 1)
        item.comment = comment
        seen_comments.add(comment_key(comment))


def apply_fallback_summaries(items: list[NewsItem], profile: dict[str, Any] | None = None) -> None:
    active_profile = profile or REPORT_PROFILES["chemistry"]
    for item in items:
        item.attractive_title = normalize_attractive_title(
            item.attractive_title or item.chinese_title,
            item,
            active_profile,
        )
        item.chinese_title = item.chinese_title or item.attractive_title
        item.comment = normalize_comment(item.comment, item)
    ensure_unique_titles_and_comments(items, active_profile)


def topic_signature(item: NewsItem, profile: dict[str, Any]) -> str:
    terms = extract_cn_terms(item, profile, limit=2)
    if terms:
        return "terms:" + "|".join(terms)
    title_key = title_fingerprint(item.title)
    return f"title:{title_key[:48]}"


def diversify_top_ids(items: list[NewsItem], top_ids: list[str], profile: dict[str, Any], desired: int = 5) -> list[str]:
    by_id = {item.item_id: item for item in items}
    selected: list[str] = []
    seen_item_titles: set[str] = set()
    seen_display_titles: set[str] = set()
    seen_topics: set[str] = set()

    def try_add(item: NewsItem, strict: bool = True) -> bool:
        if item.item_id in selected:
            return False
        original_key = title_fingerprint(item.title)
        display_key = title_body_key(display_title(item))
        topic_key = topic_signature(item, profile)
        if strict and (original_key in seen_item_titles or display_key in seen_display_titles or topic_key in seen_topics):
            return False
        selected.append(item.item_id)
        if original_key:
            seen_item_titles.add(original_key)
        if display_key:
            seen_display_titles.add(display_key)
        if topic_key:
            seen_topics.add(topic_key)
        return True

    for item_id in top_ids:
        item = by_id.get(item_id)
        if item:
            try_add(item, strict=True)
        if len(selected) >= desired:
            return selected[:desired]

    for item in items:
        try_add(item, strict=True)
        if len(selected) >= desired:
            return selected[:desired]

    for item in items:
        try_add(item, strict=False)
        if len(selected) >= desired:
            break
    return selected[:desired]


def chat_response_text(response: Any) -> str:
    choices = getattr(response, "choices", []) or []
    if not choices:
        return ""
    message = getattr(choices[0], "message", None)
    if message is None:
        return ""
    content = getattr(message, "content", "")
    if isinstance(content, list):
        pieces = []
        for part in content:
            if isinstance(part, dict):
                pieces.append(str(part.get("text") or part.get("content") or ""))
            else:
                pieces.append(str(getattr(part, "text", "") or getattr(part, "content", "")))
        return "\n".join(piece for piece in pieces if piece)
    return str(content or "")


def parse_json_object(raw: str) -> dict[str, Any]:
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", text, flags=re.DOTALL)
        if not match:
            raise
        return json.loads(match.group(0))


def resolve_llm_config(model_override: str = "") -> LLMConfig | None:
    provider = os.getenv("LLM_PROVIDER", "openai").strip().lower() or "openai"
    if provider not in SUPPORTED_LLM_PROVIDERS:
        LOGGER.warning(
            "Unsupported LLM_PROVIDER=%s; supported values are openai or deepseek. "
            "Using fallback summaries.",
            provider,
        )
        return None

    if provider == "deepseek":
        return LLMConfig(
            provider="deepseek",
            model=model_override or os.getenv("DEEPSEEK_MODEL") or DEFAULT_DEEPSEEK_MODEL,
            api_key=os.getenv("DEEPSEEK_API_KEY", ""),
            api_key_env="DEEPSEEK_API_KEY",
            base_url="https://api.deepseek.com",
        )

    return LLMConfig(
        provider="openai",
        model=model_override or os.getenv("OPENAI_MODEL") or DEFAULT_OPENAI_MODEL,
        api_key=os.getenv("OPENAI_API_KEY", ""),
        api_key_env="OPENAI_API_KEY",
    )


def generate_ai_summaries(
    items: list[NewsItem],
    model: str,
    max_ai_items: int,
    profile: dict[str, Any],
) -> dict[str, Any]:
    if not items:
        return {"top_ids": [], "field_summaries": [], "ai_generated": False}
    if OpenAI is None:
        LOGGER.warning("openai package is not installed; using fallback summaries.")
        apply_fallback_summaries(items, profile)
        return fallback_report_payload(items, profile)

    llm_config = resolve_llm_config(model)
    if llm_config is None:
        apply_fallback_summaries(items, profile)
        return fallback_report_payload(items, profile)
    if not llm_config.api_key:
        LOGGER.warning("%s is not set; using fallback summaries.", llm_config.api_key_env)
        apply_fallback_summaries(items, profile)
        return fallback_report_payload(items, profile)

    payload = [
        {
            "id": item.item_id,
            "field": item.field_name,
            "source": item.source,
            "published": format_date(item.published),
            "title": item.title,
            "source_alias": source_title_prefix(item.source, profile),
            "abstract": truncate(item.abstract, 900),
            "link": item.link,
        }
        for item in items[:max_ai_items]
    ]
    instructions = (
        f"你是{profile['ai_role']}。请基于输入论文/资讯元数据生成中文日报素材。"
        "要求准确、克制，不夸大结论；如果摘要不足，要说明信息有限。"
        f"标题风格要求：{profile.get('title_style', DEFAULT_TITLE_STYLE_GUIDE)}"
        "只输出 JSON，不要输出 Markdown。"
    )
    discipline_name = profile["title"].replace("科研资讯日报", "")
    title_schema_text = (
        "严谨的学术亮点中文标题，18-36个中文字符左右，必须以来源/期刊名开头；"
        "标题主体应突出可核对的研究对象、方法、机制、模型、数据类型、材料/体系或证据边界；"
        "不要使用反问句、悬念句、拟人化比喻、营销号语气或“不是……而是……”“一个……让……”句式"
    )
    attractive_title_prompt_text = (
        f"请根据论文信息生成一个严谨的{discipline_name}学术亮点标题。要求：1. 必须以期刊/来源名开头。"
        "2. 标题主体聚焦题名或摘要中可确认的研究对象、方法、机制、模型、数据类型、材料/体系或证据边界。"
        "3. 不使用反问、悬念、夸张、拟人化比喻和营销号句式；不使用“谁说”“不是……而是……”“一个……让……”。"
        "4. 不添加输入中没有的性能数值、应用承诺、团队、作者或因果结论。"
        "5. 可保留专业缩写、化学式、基因名、蛋白名、模型名和期刊缩写。"
    )
    single_item_prompt_template = (
        f"请根据以下论文信息生成一个严谨的{discipline_name}学术亮点标题。\n"
        "要求：\n"
        "1. 必须保留期刊/来源名作为标题开头。\n"
        "2. 标题突出研究对象、方法、机制、模型、数据类型、材料/体系或证据边界。\n"
        "3. 不使用“谁说”“不是……而是……”“一个……让……”等营销号句式。\n"
        "4. 不夸大结论，不制造虚假因果，不加入输入中没有的数值或应用承诺。\n\n"
        "论文信息：\n"
        "来源：{source}\n"
        "英文标题：{title}\n"
        "摘要：{abstract}\n"
        "领域：{field}\n\n"
        "输出只给一个标题。"
    )
    title_style_rules = [
        profile.get("title_style", DEFAULT_TITLE_STYLE_GUIDE),
        "每个 attractive_title 必须是严谨的学术亮点标题，不要写成公众号导读标题。",
        "优先把 source_alias 放在标题开头，保留期刊、预印本平台或数据库来源名称。",
        "标题主体必须体现该条目的独有关键词，例如研究对象、方法、机制、模型、数据类型、材料/体系或证据边界。",
        "禁止使用“谁说”“不是……而是……”“一个……让……”“藏在……里”“有戏”等悬念式或营销式表达。",
        "可以保留专业缩写、化学式、基因名、蛋白名、模型名和期刊缩写；不要使用“重磅”“震惊”“颠覆”等夸张词。",
        "不能编造输入中没有的团队、学校、通讯作者、性能数值、临床结论或应用承诺。",
        "不要把相关性写成因果；不能夸大临床、产业或应用价值。",
    ]
    comment_rules = [
        "comment 必须使用中文表达；可以保留必要英文缩写、化学式和物种名。",
        "不要输出以 ABSTRACT、SUMMARY 开头的英文原文片段。",
        "同批 comment 不要反复使用同一个句式；必须根据每篇题名和摘要写出不同的研究对象或证据边界。",
        "摘要不足时直接说明信息有限，并提示查看原文。",
    ]
    if profile.get("key") == "biology":
        comment_rules.extend(
            [
                "生物日报 comment 必须点出具体疾病、细胞类型、基因/蛋白/RNA、实验方法或机制对象中的至少两类。",
                "不要只写“围绕 CRISPR、转录组、发育展开”这类关键词堆砌句。",
                "CRISPR、表观遗传、临床、病毒等术语只有在英文题名或摘要明确出现时才能写入标题或摘要。",
                "结构化英文摘要要用中文整合 BACKGROUND/METHODS/RESULTS/CONCLUSION 信息，不要直接翻译或复制英文段落。",
            ]
        )
    prompt = {
        "task": profile["ai_task"],
        "schema": {
            "top_ids": ["N001"],
            "field_summaries": [
                {"field": next(iter(profile["field_keywords"])), "summary": "80字以内中文概述"}
            ],
            "items": [
                {
                    "id": "N001",
                    "attractive_title": title_schema_text,
                    "comment": "60-100字简短中文摘要，说明研究对象、方法或发现边界；不要直接复制英文摘要",
                }
            ],
        },
        "attractive_title_prompt": attractive_title_prompt_text,
        "single_item_prompt_template": single_item_prompt_template,
        "comment_rules": comment_rules,
        "title_style_rules": title_style_rules,
        "selection_rules": [
            "top_ids 选 5 条最值得关注的条目，兼顾来源权威性、新近性和领域覆盖。",
            "items 必须覆盖输入中的每一个 id。",
            "field_summaries 只覆盖输入中实际出现的领域。",
        ],
        "strict_json_rules": [
            "输出必须是一个可被 json.loads 直接解析的 JSON object。",
            "不要在 JSON 前后添加解释、Markdown 代码块或注释。",
            "所有字符串里的换行、引号和反斜杠都必须正确转义。",
        ],
        "items": payload,
    }

    client_kwargs = {"api_key": llm_config.api_key}
    if llm_config.base_url:
        client_kwargs["base_url"] = llm_config.base_url
    client = OpenAI(**client_kwargs)

    def request_ai_json(user_prompt_base: dict[str, Any], max_tokens: int, label: str) -> dict[str, Any]:
        local_last_error: Exception | None = None
        for attempt in range(2):
            user_prompt = user_prompt_base
            if attempt:
                user_prompt = {
                    **user_prompt_base,
                    "retry_instruction": "上一次输出不是合法 JSON。请只返回严格 JSON object，不要省略逗号，不要输出 Markdown。",
                }
            request_kwargs = {
                "model": llm_config.model,
                "messages": [
                    {"role": "system", "content": instructions},
                    {"role": "user", "content": json.dumps(user_prompt, ensure_ascii=False)},
                ],
                "temperature": 0.2,
                "max_tokens": max_tokens,
                "response_format": {"type": "json_object"},
            }
            try:
                try:
                    response = client.chat.completions.create(**request_kwargs)
                except Exception as exc:  # noqa: BLE001 - some compatible providers reject response_format.
                    local_last_error = exc
                    request_kwargs.pop("response_format", None)
                    response = client.chat.completions.create(**request_kwargs)
                raw_response = chat_response_text(response)
                return parse_json_object(raw_response)
            except Exception as exc:  # noqa: BLE001 - AI failure should not block the document.
                local_last_error = exc
                LOGGER.warning("%s summary generation %s attempt %d failed: %s", llm_config.provider, label, attempt + 1, exc)
        raise RuntimeError(str(local_last_error or "unknown AI JSON error"))

    def request_ai_json_in_chunks(chunk_size: int = 10) -> dict[str, Any]:
        combined_items: list[dict[str, Any]] = []
        combined_top_ids: list[str] = []
        field_summary_by_name: dict[str, str] = {}
        for chunk_index, start in enumerate(range(0, len(payload), chunk_size), start=1):
            chunk_payload = payload[start : start + chunk_size]
            chunk_prompt = {
                **prompt,
                "items": chunk_payload,
                "selection_rules": [
                    "top_ids 从本批条目中选 1-3 条最值得关注的条目。",
                    "items 必须覆盖本批输入中的每一个 id。",
                    "field_summaries 只覆盖本批输入中实际出现的领域。",
                ],
            }
            chunk_parsed = request_ai_json(chunk_prompt, 4500, f"chunk {chunk_index}")
            parsed_items = chunk_parsed.get("items", [])
            if isinstance(parsed_items, list):
                combined_items.extend(entry for entry in parsed_items if isinstance(entry, dict))
            for item_id in chunk_parsed.get("top_ids", []):
                if isinstance(item_id, str) and item_id not in combined_top_ids:
                    combined_top_ids.append(item_id)
            summaries = chunk_parsed.get("field_summaries", [])
            if isinstance(summaries, list):
                for summary in summaries:
                    if not isinstance(summary, dict):
                        continue
                    field_name = clean_text(summary.get("field"))
                    summary_text = clean_text(summary.get("summary"))
                    if field_name and summary_text and field_name not in field_summary_by_name:
                        field_summary_by_name[field_name] = summary_text
        return {
            "top_ids": combined_top_ids,
            "field_summaries": [
                {"field": field_name, "summary": summary}
                for field_name, summary in field_summary_by_name.items()
            ],
            "items": combined_items,
        }

    parsed: dict[str, Any] | None = None
    last_error: Exception | None = None
    try:
        if llm_config.provider == "deepseek" and len(payload) > 18:
            LOGGER.info("Using chunked %s summary generation for %d items.", llm_config.provider, len(payload))
            parsed = request_ai_json_in_chunks()
        else:
            parsed = request_ai_json(prompt, 9000, "full")
    except Exception as exc:  # noqa: BLE001 - AI failure should not block the document.
        last_error = exc
        LOGGER.warning("%s summary generation failed: %s", llm_config.provider, exc)

    if parsed is None:
        LOGGER.warning("%s summary generation failed; using fallback summaries: %s", llm_config.provider, last_error)
        apply_fallback_summaries(items, profile)
        return fallback_report_payload(items, profile)

    by_id = {entry.get("id"): entry for entry in parsed.get("items", []) if isinstance(entry, dict)}
    complete_ai_items = True
    missing_ai_item_ids: list[str] = []
    for item in items:
        generated = by_id.get(item.item_id, {})
        generated_title = generated.get("attractive_title", "") or generated.get("chinese_title", "")
        generated_comment = generated.get("comment", "")
        if not clean_text(generated_title) or not clean_text(generated_comment):
            complete_ai_items = False
            missing_ai_item_ids.append(item.item_id)
        item.attractive_title = normalize_attractive_title(generated_title, item, profile)
        item.chinese_title = item.chinese_title or item.attractive_title
        item.comment = normalize_comment(generated_comment, item)
    if missing_ai_item_ids:
        LOGGER.warning(
            "%s summary generation did not return title/comment for item ids: %s",
            llm_config.provider,
            ", ".join(missing_ai_item_ids[:10]),
        )
    apply_fallback_summaries(items, profile)
    top_ids = [
        item_id
        for item_id in parsed.get("top_ids", [])
        if isinstance(item_id, str) and any(item.item_id == item_id for item in items)
    ]
    if len(top_ids) < 5:
        existing = set(top_ids)
        for item in items:
            if item.item_id not in existing:
                top_ids.append(item.item_id)
            if len(top_ids) >= 5:
                break
    top_ids = diversify_top_ids(items, top_ids, profile, desired=5)
    field_summaries = parsed.get("field_summaries", [])
    if not isinstance(field_summaries, list):
        field_summaries = []
    return {
        "top_ids": top_ids[:5],
        "field_summaries": field_summaries,
        "ai_generated": complete_ai_items,
        "ai_provider": llm_config.provider,
        "ai_model": llm_config.model,
    }


def fallback_report_payload(items: list[NewsItem], profile: dict[str, Any] | None = None) -> dict[str, Any]:
    active_profile = profile or REPORT_PROFILES["chemistry"]
    top_ids = diversify_top_ids(items, [item.item_id for item in items[:5]], active_profile, desired=5)
    grouped: dict[str, list[NewsItem]] = {}
    for item in items:
        grouped.setdefault(item.field_name, []).append(item)
    field_summaries = [
        {
            "field": field_name,
            "summary": f"共检索到 {len(group_items)} 条相关资讯，主要来源包括："
            f"{'、'.join(sorted({entry.source for entry in group_items})[:4])}。",
        }
        for field_name, group_items in grouped.items()
    ]
    return {"top_ids": top_ids, "field_summaries": field_summaries, "ai_generated": False}


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    relationship_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)  # noqa: SLF001 - python-docx hyperlink helper requires OXML.


def set_run_font(
    run: Any,
    name: str = "Arial",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)  # noqa: SLF001
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)  # noqa: SLF001
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)  # noqa: SLF001
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph: Any, color: str = "D9E2EC", size: str = "8") -> None:
    properties = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)


def set_paragraph_shading(paragraph: Any, fill: str = "F6F8FA") -> None:
    properties = paragraph._p.get_or_add_pPr()  # noqa: SLF001
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def add_spacer(document: Document, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)


def set_document_fonts(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)

    styles = document.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")  # noqa: SLF001
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")  # noqa: SLF001
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")  # noqa: SLF001
    normal = styles["Normal"]
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.28

    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor(17, 24, 39)
    styles["Title"].paragraph_format.space_after = Pt(4)

    styles["Heading 1"].font.size = Pt(13)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(15, 76, 117)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)

    styles["Heading 2"].font.size = Pt(11.5)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(15, 76, 117)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    styles["Heading 3"].font.size = Pt(10.3)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor(17, 24, 39)
    styles["Heading 3"].paragraph_format.space_before = Pt(6)
    styles["Heading 3"].paragraph_format.space_after = Pt(2)


def add_label_value(document: Document, label: str, value: str, link: str = "") -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}：")
    label_run.bold = True
    if link:
        add_hyperlink(paragraph, value, link)
    else:
        paragraph.add_run(value)


def add_masthead(document: Document, report_date: date, item_count: int, profile: dict[str, Any]) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(profile["header_label"])
    set_run_font(run, size=8.5, color=RGBColor(107, 114, 128), bold=True)
    paragraph_border_bottom(header, color="E5E7EB", size="4")

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Generated by science-news-daily")
    set_run_font(run, size=8, color=RGBColor(156, 163, 175))

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("DAILY RESEARCH BRIEF")
    set_run_font(run, size=8.5, color=RGBColor(15, 76, 117), bold=True)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(7)
    run = title.add_run(profile["title"])
    set_run_font(run, size=22, color=RGBColor(17, 24, 39), bold=True)

    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(2)
    meta.paragraph_format.space_after = Pt(11)
    meta.paragraph_format.line_spacing = 1.2
    run = meta.add_run(f"{report_date.isoformat()}  /  精选 {item_count} 篇  /  {profile['meta_fields']}")
    set_run_font(run, size=9.2, color=RGBColor(75, 85, 99))
    paragraph_border_bottom(meta, color="CBD5E1", size="10")


def add_source_note(document: Document, item_count: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(9)
    run = paragraph.add_run(
        f"本期收录 {item_count} 条，经来源权重、研究新近性、摘要信息量和学习价值信号排序；单源异常会在运行诊断中标记。"
    )
    set_run_font(run, size=9, color=RGBColor(75, 85, 99))


def add_top_item_block(document: Document, item: NewsItem, index: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(1)
    label = paragraph.add_run(f"{index:02d}  ")
    set_run_font(label, size=9.2, color=RGBColor(15, 76, 117), bold=True)
    title_run = paragraph.add_run(display_title(item))
    set_run_font(title_run, size=10.2, color=RGBColor(17, 24, 39), bold=True)

    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(1)
    run = meta.add_run(f"{item.field_name}  |  {item.source}  |  {format_date(item.published)}")
    set_run_font(run, size=8.5, color=RGBColor(107, 114, 128))

    comment = document.add_paragraph()
    comment.paragraph_format.space_before = Pt(0)
    comment.paragraph_format.space_after = Pt(6)
    comment.paragraph_format.line_spacing = 1.22
    run = comment.add_run(truncate(item.comment or fallback_comment(item), 150))
    set_run_font(run, size=9, color=RGBColor(55, 65, 81))
    paragraph_border_bottom(comment, color="EEF2F7", size="4")


def add_item_block(document: Document, item: NewsItem) -> None:
    heading = document.add_paragraph()
    heading.style = document.styles["Heading 3"]
    heading.paragraph_format.keep_with_next = True
    heading.add_run(display_title(item))

    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    run = meta.add_run(f"{item.source}  |  {format_date(item.published)}  |  {item.field_name}")
    set_run_font(run, size=8.5, color=RGBColor(107, 114, 128))

    english = document.add_paragraph()
    english.paragraph_format.space_before = Pt(0)
    english.paragraph_format.space_after = Pt(2)
    label = english.add_run("英文标题  ")
    set_run_font(label, size=8.2, color=RGBColor(15, 76, 117), bold=True)
    run = english.add_run(item.title)
    set_run_font(run, size=9, color=RGBColor(55, 65, 81), italic=True)

    comment = document.add_paragraph()
    comment.paragraph_format.space_before = Pt(1)
    comment.paragraph_format.space_after = Pt(4)
    comment.paragraph_format.line_spacing = 1.24
    comment.paragraph_format.left_indent = Inches(0.06)
    set_paragraph_shading(comment, fill="F8FAFC")
    label = comment.add_run("中文摘要  ")
    set_run_font(label, size=8.7, color=RGBColor(15, 76, 117), bold=True)
    run = comment.add_run(item.comment or fallback_comment(item))
    set_run_font(run, size=9.2, color=RGBColor(31, 41, 55))

    abstract = truncate(item.abstract, 420)
    if abstract:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.22
        label = paragraph.add_run("原文摘要  ")
        set_run_font(label, size=8.7, color=RGBColor(107, 114, 128), bold=True)
        run = paragraph.add_run(abstract)
        set_run_font(run, size=8.8, color=RGBColor(75, 85, 99))

    links = document.add_paragraph()
    links.paragraph_format.space_before = Pt(0)
    links.paragraph_format.space_after = Pt(7)
    links.paragraph_format.line_spacing = 1.18
    label = links.add_run("链接  ")
    set_run_font(label, size=8.5, color=RGBColor(107, 114, 128), bold=True)
    add_hyperlink(links, item.link, item.link)
    if item.doi:
        links.add_run("   DOI  ")
        add_hyperlink(links, item.doi, f"https://doi.org/{item.doi}")
    paragraph_border_bottom(links, color="E5E7EB", size="4")


def diagnostics_has_issue(diagnostics: Any | None) -> bool:
    return bool(diagnostics is not None and not getattr(diagnostics, "network_ok", False))


def source_has_failures(source_statuses: list[SourceStatus] | None) -> bool:
    return any(not status.success for status in source_statuses or [])


def add_run_diagnostics_section(
    document: Document,
    diagnostics: Any | None,
    source_statuses: list[SourceStatus] | None,
) -> None:
    if not diagnostics_has_issue(diagnostics) and not source_has_failures(source_statuses):
        return

    document.add_heading("运行诊断", level=1)
    if diagnostics is not None:
        if getattr(diagnostics, "network_ok", False):
            document.add_paragraph("网络诊断：DNS 与 HTTPS 探测均正常。")
        else:
            document.add_paragraph("网络诊断：发现 DNS 或 HTTPS 访问异常。")
        for line in diagnostics.summary_lines():
            document.add_paragraph(line, style=document.styles["List Bullet"])
    else:
        document.add_paragraph("网络诊断：未执行。")

    failed_statuses = [status for status in source_statuses or [] if not status.success]
    if failed_statuses:
        document.add_paragraph("失败来源：")
        for status in failed_statuses:
            document.add_paragraph(
                f"{status.name}: {summarize_source_error(status.error)}",
                style=document.styles["List Bullet"],
            )


def create_document(
    items: list[NewsItem],
    report_payload: dict[str, Any],
    report_date: date,
    output_dir: Path,
    profile: dict[str, Any],
    diagnostics: Any | None = None,
    source_statuses: list[SourceStatus] | None = None,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_fonts(document)

    add_masthead(document, report_date, len(items), profile)
    add_source_note(document, len(items))

    by_id = {item.item_id: item for item in items}
    top_items = [by_id[item_id] for item_id in report_payload.get("top_ids", []) if item_id in by_id]
    if len(top_items) < 5:
        existing = {item.item_id for item in top_items}
        top_items.extend([item for item in items if item.item_id not in existing][: 5 - len(top_items)])

    document.add_heading("今日重点 5 条", level=1)
    for index, item in enumerate(top_items[:5], start=1):
        add_top_item_block(document, item, index)

    document.add_heading("分领域摘要", level=1)
    field_summary_map: dict[str, str] = {}
    for entry in report_payload.get("field_summaries", []):
        if isinstance(entry, dict) and entry.get("field") and entry.get("summary"):
            field_summary_map[clean_text(entry["field"])] = clean_text(entry["summary"])

    grouped: dict[str, list[NewsItem]] = {}
    for item in items:
        grouped.setdefault(item.field_name, []).append(item)

    ordered_fields = list(profile["field_keywords"].keys()) + [profile["default_field"]]
    for field_name in ordered_fields:
        group_items = grouped.get(field_name, [])
        if not group_items:
            continue
        document.add_heading(field_name, level=2)
        summary = field_summary_map.get(
            field_name,
            f"本领域收录 {len(group_items)} 条，主要覆盖近期论文、期刊上线内容和科研资讯。",
        )
        summary_paragraph = document.add_paragraph()
        summary_paragraph.paragraph_format.space_after = Pt(6)
        summary_paragraph.paragraph_format.line_spacing = 1.25
        run = summary_paragraph.add_run(summary)
        set_run_font(run, size=9.2, color=RGBColor(55, 65, 81))
        for item in group_items:
            add_item_block(document, item)

    add_run_diagnostics_section(document, diagnostics, source_statuses)

    output_path = output_dir / f"{profile['output_prefix']}_{report_date.isoformat()}.docx"
    document.save(output_path)
    return output_path


def source_failure_recommendations(
    diagnostics: Any | None,
    source_statuses: list[SourceStatus],
    zero_items: bool,
) -> list[str]:
    recommendations: list[str] = []
    dns_failed_hosts = getattr(diagnostics, "dns_failed_hosts", []) if diagnostics else []
    https_failed_hosts = getattr(diagnostics, "https_failed_hosts", []) if diagnostics else []
    failed_statuses = [status for status in source_statuses if not status.success]

    if dns_failed_hosts:
        recommendations.append(
            "检查 DNS、代理或 VPN 配置，确认 arxiv.org、pubmed.ncbi.nlm.nih.gov、api.crossref.org 可以解析。"
        )
    if https_failed_hosts:
        recommendations.append(
            "检查 HTTPS 出口、代理证书、公司/校园网 TLS 拦截，以及 HTTP_PROXY/HTTPS_PROXY/ALL_PROXY 环境变量。"
        )
    if failed_statuses:
        recommendations.append(
            "查看日志中失败来源的异常信息；单个来源失败通常可以稍后重试，或暂时降低 --source-limit。"
        )
    if zero_items and not failed_statuses:
        recommendations.append(
            "网络和来源没有抛出异常但结果为 0 条时，可扩大 --days、提高 --source-limit，或稍后重试。"
        )
    recommendations.append("在项目目录运行 `.venv/bin/python network_check.py` 可单独复查网络。")
    recommendations.append("手动指定输出目录时使用 `--output-dir /path/to/output`，默认输出在项目内 `./output`。")
    return recommendations


def create_failure_report(
    report_date: date,
    output_dir: Path,
    profile: dict[str, Any],
    diagnostics: Any | None,
    source_statuses: list[SourceStatus],
    reason: str,
    collected_count: int,
    prepared_count: int,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_fonts(document)

    title = document.add_heading(profile["failure_title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(192, 0, 0)

    document.add_paragraph(f"日期：{report_date.isoformat()}")
    document.add_paragraph(f"失败原因：{reason}")
    document.add_paragraph(f"原始抓取条数：{collected_count}；去重过滤后条数：{prepared_count}")

    document.add_heading("DNS 是否失败", level=1)
    if diagnostics is None:
        document.add_paragraph("未执行网络诊断。")
    else:
        dns_failed_hosts = getattr(diagnostics, "dns_failed_hosts", [])
        if dns_failed_hosts:
            document.add_paragraph(f"DNS 失败主机：{'、'.join(dns_failed_hosts)}")
        else:
            document.add_paragraph("DNS 解析未发现失败。")
        document.add_heading("网络诊断详情", level=2)
        for line in diagnostics.summary_lines():
            document.add_paragraph(line, style=document.styles["List Bullet"])

    document.add_heading("哪些来源失败", level=1)
    if source_statuses:
        for status in source_statuses:
            document.add_paragraph(f"{status.name}: {status.summary()}", style=document.styles["List Bullet"])
    else:
        document.add_paragraph("没有记录到来源状态；可能在初始化阶段失败。")

    document.add_heading("建议修复动作", level=1)
    for recommendation in source_failure_recommendations(diagnostics, source_statuses, True):
        document.add_paragraph(recommendation, style=document.styles["List Bullet"])

    failure_filename = "运行失败报告.docx"
    if profile["key"] != "chemistry":
        failure_filename = f"{profile['output_prefix']}_运行失败报告.docx"
    output_path = output_dir / failure_filename
    document.save(output_path)
    return output_path


def all_sources_failed(source_statuses: list[SourceStatus]) -> bool:
    return bool(source_statuses) and all(not status.success for status in source_statuses)


def resolve_output_dir(value: str) -> Path:
    path = Path(value).expanduser()
    if path.is_absolute():
        return path
    return Path(__file__).resolve().parent / path


def env_flag(name: str, default: bool = True) -> bool:
    value = os.getenv(name)
    if value is None:
        return default
    return value.strip().lower() not in {"0", "false", "no", "off"}


def parse_email_recipients(value: str) -> list[str]:
    normalized = re.sub(r"[;\r\n]+", ",", value)
    recipients: list[str] = []
    seen: set[str] = set()
    for _, address in getaddresses([normalized]):
        address = address.strip()
        if not address or "@" not in address:
            continue
        key = address.lower()
        if key in seen:
            continue
        recipients.append(address)
        seen.add(key)
    return recipients


def find_libreoffice_executable() -> str | None:
    configured = os.getenv("LIBREOFFICE_PATH", "").strip()
    candidates: list[str] = []
    if configured:
        candidates.append(configured)
        resolved = shutil.which(configured)
        if resolved:
            candidates.append(resolved)

    for command in ("soffice", "libreoffice"):
        resolved = shutil.which(command)
        if resolved:
            candidates.append(resolved)

    candidates.extend(
        [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/usr/local/bin/soffice",
            "/usr/bin/soffice",
            "/usr/bin/libreoffice",
            "/snap/bin/libreoffice",
        ]
    )

    seen: set[str] = set()
    for candidate in candidates:
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        candidate_path = Path(candidate)
        if candidate_path.exists() and os.access(candidate_path, os.X_OK):
            return str(candidate_path)
    return None


def convert_docx_to_pdf(docx_path: Path) -> Path | None:
    if not docx_path.exists():
        LOGGER.warning("PDF conversion skipped; DOCX does not exist: %s", docx_path)
        return None
    if docx_path.suffix.lower() != ".docx":
        LOGGER.warning("PDF conversion skipped; input is not a DOCX file: %s", docx_path)
        return None

    soffice = find_libreoffice_executable()
    if not soffice:
        LOGGER.warning(
            "PDF conversion skipped; LibreOffice executable was not found. "
            "Install LibreOffice or set LIBREOFFICE_PATH."
        )
        return None

    pdf_path = docx_path.with_suffix(".pdf")
    timeout_raw = os.getenv("PDF_CONVERT_TIMEOUT", "120").strip()
    try:
        timeout = max(10, int(timeout_raw))
    except ValueError:
        timeout = 120

    started_at = time.time()
    with tempfile.TemporaryDirectory(prefix="chem-news-lo-") as profile_dir:
        command = [
            soffice,
            f"-env:UserInstallation={Path(profile_dir).as_uri()}",
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(docx_path.parent),
            str(docx_path),
        ]
        try:
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=timeout,
                check=False,
            )
        except subprocess.TimeoutExpired:
            LOGGER.warning("PDF conversion timed out after %s seconds: %s", timeout, docx_path)
            return None
        except Exception as exc:  # noqa: BLE001 - conversion failure should not invalidate DOCX output.
            LOGGER.warning("PDF conversion failed to start: %s", exc)
            return None

    if result.returncode != 0:
        LOGGER.warning(
            "PDF conversion failed with exit code %s. stdout=%s stderr=%s",
            result.returncode,
            result.stdout.strip(),
            result.stderr.strip(),
        )
        return None

    if not pdf_path.exists() or pdf_path.stat().st_size == 0:
        LOGGER.warning("PDF conversion did not produce a valid file: %s", pdf_path)
        return None
    if pdf_path.stat().st_mtime < started_at - 2:
        LOGGER.warning("PDF conversion output appears stale and will not be sent: %s", pdf_path)
        return None

    LOGGER.info("Converted DOCX to PDF: %s", pdf_path)
    return pdf_path


def send_report_email(
    attachment_path: Path,
    report_date: date,
    profile: dict[str, Any],
    is_failure: bool = False,
    ai_generated: bool = False,
) -> bool:
    if not env_flag("EMAIL_ENABLED", True):
        LOGGER.info("Email sending is disabled by EMAIL_ENABLED.")
        return False

    smtp_host = os.getenv("SMTP_HOST", "").strip()
    smtp_username = os.getenv("SMTP_USERNAME", "").strip()
    smtp_password = os.getenv("SMTP_PASSWORD", "")
    smtp_from = os.getenv("SMTP_FROM", "").strip() or smtp_username
    smtp_security = os.getenv("SMTP_SECURITY", "").strip().lower() or "ssl"
    smtp_port_raw = os.getenv("SMTP_PORT", "").strip()
    smtp_port = int(smtp_port_raw) if smtp_port_raw else (587 if smtp_security in {"starttls", "tls"} else 465)
    fallback_recipient_value = os.getenv("REPORT_EMAIL_TO", "").strip() if profile["key"] == "chemistry" else ""
    recipient_value = os.getenv(profile["email_env"], "").strip() or fallback_recipient_value or profile["default_email_to"]
    recipients = parse_email_recipients(recipient_value)

    missing = [
        name
        for name, value in {
            "SMTP_HOST": smtp_host,
            "SMTP_USERNAME": smtp_username,
            "SMTP_PASSWORD": smtp_password,
            "SMTP_FROM or SMTP_USERNAME": smtp_from,
            profile["email_env"] if profile["key"] != "chemistry" else "CHEM_REPORT_EMAIL_TO or REPORT_EMAIL_TO": ",".join(recipients),
        }.items()
        if not value
    ]
    if missing:
        LOGGER.warning("Email not sent; missing SMTP config: %s", ", ".join(missing))
        return False
    if not attachment_path.exists():
        LOGGER.warning("Email not sent; attachment does not exist: %s", attachment_path)
        return False
    if not is_failure and not ai_generated:
        LOGGER.warning("Email not sent; normal report does not have complete AI-generated summaries.")
        return False

    email_attachment_path = attachment_path
    local_docx_name = ""
    if attachment_path.suffix.lower() == ".docx":
        local_docx_name = attachment_path.name
        converted_path = convert_docx_to_pdf(attachment_path)
        if converted_path is None:
            LOGGER.warning("Email not sent; PDF conversion failed for %s", attachment_path)
            return False
        email_attachment_path = converted_path
    elif attachment_path.suffix.lower() != ".pdf":
        LOGGER.warning("Email not sent; only DOCX-to-PDF or PDF attachments are supported: %s", attachment_path)
        return False

    subject_prefix = f"{profile['title']}运行失败" if is_failure else profile["title"]
    message = EmailMessage()
    message["Subject"] = f"{subject_prefix} - {report_date.isoformat()}"
    message["From"] = smtp_from
    message["To"] = ", ".join(recipients)
    body_lines = [
        f"日期：{report_date.isoformat()}",
        f"PDF附件：{email_attachment_path.name}",
    ]
    if local_docx_name:
        body_lines.append(f"本地DOCX文件：{local_docx_name}")
    body_lines.extend(["", "本邮件由 science-news-daily 自动发送。"])
    message.set_content("\n".join(body_lines))
    message.add_attachment(
        email_attachment_path.read_bytes(),
        maintype="application",
        subtype="pdf",
        filename=email_attachment_path.name,
    )

    try:
        if smtp_security == "ssl":
            with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=30) as smtp:
                smtp.login(smtp_username, smtp_password)
                smtp.send_message(message)
        else:
            with smtplib.SMTP(smtp_host, smtp_port, timeout=30) as smtp:
                if smtp_security in {"starttls", "tls"}:
                    smtp.starttls()
                smtp.login(smtp_username, smtp_password)
                smtp.send_message(message)
    except Exception as exc:  # noqa: BLE001 - email failure should not invalidate the report.
        LOGGER.warning("Email sending failed: %s", exc)
        return False

    LOGGER.info("Sent report email to %s with PDF attachment %s", ", ".join(recipients), email_attachment_path)
    return True


def collect_items(
    args: argparse.Namespace,
    since: datetime,
    until: datetime,
    profile: dict[str, Any],
) -> tuple[list[NewsItem], list[SourceStatus]]:
    session = build_session()
    fetchers: list[tuple[str, Callable[[], list[NewsItem]]]] = [
        ("arXiv", lambda: fetch_arxiv(session, since, until, args.source_limit, profile)),
        ("PubMed", lambda: fetch_pubmed(session, since, until, args.source_limit, profile)),
    ]

    all_items: list[NewsItem] = []
    statuses: list[SourceStatus] = []
    for source_name, fetcher in fetchers:
        try:
            items = fetcher()
            LOGGER.info("%s returned %d items", source_name, len(items))
            all_items.extend(items)
            statuses.append(SourceStatus(name=source_name, success=True, item_count=len(items)))
        except Exception as exc:  # noqa: BLE001 - source isolation is required.
            LOGGER.exception("%s failed and was skipped: %s", source_name, exc)
            statuses.append(
                SourceStatus(
                    name=source_name,
                    success=False,
                    error=f"{type(exc).__name__}: {exc}",
                )
            )

    try:
        crossref_items, crossref_statuses = fetch_crossref(session, since, until, args.source_limit, profile)
        LOGGER.info("Crossref returned %d items", len(crossref_items))
        all_items.extend(crossref_items)
        statuses.extend(crossref_statuses)
    except Exception as exc:  # noqa: BLE001 - defensive guard around the grouped source.
        LOGGER.exception("Crossref failed and was skipped: %s", exc)
        statuses.append(
            SourceStatus(
                name="Crossref",
                success=False,
                error=f"{type(exc).__name__}: {exc}",
            )
        )

    try:
        rss_items, rss_statuses = fetch_rss(session, since, until, args.source_limit, profile)
        LOGGER.info("RSS returned %d items", len(rss_items))
        all_items.extend(rss_items)
        statuses.extend(rss_statuses)
    except Exception as exc:  # noqa: BLE001 - defensive guard around the grouped source.
        LOGGER.exception("RSS failed and was skipped: %s", exc)
        statuses.append(
            SourceStatus(
                name="RSS",
                success=False,
                error=f"{type(exc).__name__}: {exc}",
            )
        )
    return all_items, statuses


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Generate a Chinese research news daily DOCX.")
    parser.add_argument(
        "--profile",
        choices=sorted(REPORT_PROFILES),
        default=os.getenv("REPORT_PROFILE", "chemistry"),
        help="Report profile to run: chemistry, biology, or statistics. Default: chemistry.",
    )
    parser.add_argument(
        "--days",
        type=int,
        default=int(os.getenv("CHEM_NEWS_DAYS", "3")),
        help="Lookback window in days. Use 1 for the most recent 24 hours, or 3 for a broader daily digest.",
    )
    parser.add_argument(
        "--max-items",
        type=int,
        default=int(os.getenv("CHEM_NEWS_MAX_ITEMS", str(DEFAULT_MAX_ITEMS))),
        help="Maximum number of deduplicated items written to the report.",
    )
    parser.add_argument(
        "--source-limit",
        type=int,
        default=int(os.getenv("CHEM_NEWS_SOURCE_LIMIT", "80")),
        help="Maximum records requested from each API source before filtering.",
    )
    parser.add_argument(
        "--max-ai-items",
        type=int,
        default=int(os.getenv("CHEM_NEWS_MAX_AI_ITEMS", str(DEFAULT_MAX_AI_ITEMS))),
        help="Maximum records sent to the LLM provider for Chinese title/comment generation.",
    )
    parser.add_argument(
        "--model",
        default="",
        help=(
            "Override the LLM model for this run. Defaults to OPENAI_MODEL/"
            f"{DEFAULT_OPENAI_MODEL} for openai, or DEEPSEEK_MODEL/"
            f"{DEFAULT_DEEPSEEK_MODEL} for deepseek."
        ),
    )
    parser.add_argument(
        "--output-dir",
        default=os.getenv("CHEM_NEWS_OUTPUT_DIR", DEFAULT_OUTPUT_DIR),
        help=f"Output directory. Default: {DEFAULT_OUTPUT_DIR}",
    )
    parser.add_argument(
        "--report-date",
        default=os.getenv("CHEM_NEWS_REPORT_DATE", ""),
        help="Report date in YYYY-MM-DD format. Defaults to today's local date.",
    )
    parser.add_argument(
        "--no-openai",
        action="store_true",
        help="Skip model API calls and create the DOCX with deterministic fallback comments.",
    )
    parser.add_argument(
        "--no-email",
        action="store_true",
        help="Skip SMTP delivery for this run while still generating local DOCX/PDF outputs.",
    )
    parser.add_argument(
        "--require-email",
        action="store_true",
        default=env_flag("REQUIRE_EMAIL_SUCCESS", False),
        help="Return a non-zero exit code if SMTP delivery is disabled or fails.",
    )
    parser.add_argument(
        "--require-ai",
        action="store_true",
        default=env_flag("REQUIRE_AI_SUMMARY", False),
        help="Return a non-zero exit code unless every report item received an AI-generated title and summary.",
    )
    parser.add_argument("--verbose", action="store_true", help="Enable debug logging.")
    return parser


def ensure_runtime_dependencies() -> None:
    if not MISSING_DEPENDENCIES:
        return
    missing = ", ".join(sorted(set(MISSING_DEPENDENCIES)))
    raise SystemExit(
        "Missing Python dependencies: "
        f"{missing}. Run `pip install -r requirements.txt` in this project environment."
    )


def configure_logging(verbose: bool) -> None:
    logging.basicConfig(
        level=logging.DEBUG if verbose else logging.INFO,
        format="%(asctime)s %(levelname)s %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )


def parse_report_date(value: str) -> date:
    if not value:
        return datetime.now().astimezone().date()
    try:
        return date.fromisoformat(value)
    except ValueError as exc:
        raise SystemExit(f"--report-date must use YYYY-MM-DD format: {value}") from exc


def stable_hash(value: str) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()[:10]


def ensure_item_ids(items: list[NewsItem]) -> None:
    used: set[str] = set()
    for index, item in enumerate(items, start=1):
        if item.item_id:
            continue
        base = f"N{index:03d}"
        if base not in used:
            item.item_id = base
        else:
            item.item_id = f"N{stable_hash(item.title)}"
        used.add(item.item_id)


def main() -> int:
    parser = build_arg_parser()
    args = parser.parse_args()
    configure_logging(args.verbose)
    ensure_runtime_dependencies()
    profile = resolve_profile(args.profile)
    if args.no_email and args.require_email:
        raise SystemExit("--no-email cannot be used together with --require-email.")
    if args.no_openai and args.require_ai:
        raise SystemExit("--no-openai cannot be used together with --require-ai.")
    if args.no_email:
        os.environ["EMAIL_ENABLED"] = "false"

    if args.days < 1 or args.days > 14:
        raise SystemExit("--days should be between 1 and 14.")
    if args.max_items < 1:
        raise SystemExit("--max-items must be positive.")

    now = datetime.now(timezone.utc)
    since = now - timedelta(days=args.days)
    report_date = parse_report_date(args.report_date)
    output_dir = resolve_output_dir(args.output_dir)

    from network_check import run_network_checks

    diagnostics = run_network_checks(logger=LOGGER)
    if not diagnostics.network_ok:
        LOGGER.warning("Network is unavailable or degraded: %s", " | ".join(diagnostics.summary_lines()))

    LOGGER.info(
        "Collecting %s items from %s to %s",
        profile["key"],
        since.isoformat(),
        now.isoformat(),
    )
    collected, source_statuses = collect_items(args, since, now, profile)
    prepared = prepare_items(collected, args.max_items, now, profile)
    ensure_item_ids(prepared)
    LOGGER.info("Prepared %d deduplicated items", len(prepared))

    if not prepared:
        reason = "抓取和过滤后没有可写入日报的资讯。"
        if all_sources_failed(source_statuses):
            reason = "全部来源抓取失败，未获得任何资讯。"
        failure_report = create_failure_report(
            report_date=report_date,
            output_dir=output_dir,
            profile=profile,
            diagnostics=diagnostics,
            source_statuses=source_statuses,
            reason=reason,
            collected_count=len(collected),
            prepared_count=len(prepared),
        )
        LOGGER.error("No reportable items; saved failure report to %s", failure_report)
        if args.require_ai:
            LOGGER.error("AI summary is required; failure report email will not be sent.")
            print(failure_report)
            return 4
        email_sent = send_report_email(failure_report, report_date, profile, is_failure=True)
        print(failure_report)
        if args.require_email and not email_sent:
            return 3
        return 2 if all_sources_failed(source_statuses) else 1

    if args.no_openai:
        apply_fallback_summaries(prepared, profile)
        report_payload = fallback_report_payload(prepared, profile)
    else:
        if args.require_ai and args.max_ai_items < len(prepared):
            LOGGER.error(
                "AI summary is required for every emailed item, but --max-ai-items=%d is less than prepared item count=%d.",
                args.max_ai_items,
                len(prepared),
            )
            return 4
        report_payload = generate_ai_summaries(prepared, args.model, args.max_ai_items, profile)

    if args.require_ai and not report_payload.get("ai_generated"):
        LOGGER.error(
            "AI summary is required, but model generation was incomplete or fell back to rule-based summaries. "
            "Report email will not be sent."
        )
        return 4

    output_path = create_document(
        prepared,
        report_payload,
        report_date,
        output_dir,
        profile,
        diagnostics=diagnostics,
        source_statuses=source_statuses,
    )
    LOGGER.info("Saved report to %s", output_path)
    email_sent = send_report_email(output_path, report_date, profile, ai_generated=bool(report_payload.get("ai_generated")))
    print(output_path)
    if args.require_email and not email_sent:
        return 3
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except KeyboardInterrupt:
        print("Interrupted.", file=sys.stderr)
        raise SystemExit(130)
