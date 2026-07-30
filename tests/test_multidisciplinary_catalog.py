import os
import unittest
from datetime import UTC, datetime
from datetime import date
from pathlib import Path
from tempfile import TemporaryDirectory
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

import main


FIRST_LEVEL_PROFILE_IDS = {
    "philosophy",
    "economics",
    "law",
    "education",
    "literature",
    "history",
    "natural_sciences",
    "engineering",
    "agriculture",
    "medicine",
    "management",
    "arts",
    "interdisciplinary_studies",
    "military_science",
}


class MultidisciplinaryCatalogueTests(unittest.TestCase):
    def test_all_first_level_disciplines_and_computer_science_are_available(self) -> None:
        self.assertTrue(FIRST_LEVEL_PROFILE_IDS.issubset(main.REPORT_PROFILES))
        self.assertIn("computer_science", main.REPORT_PROFILES)
        self.assertEqual(main.resolve_profile("chemistry")["key"], "chemistry")
        self.assertEqual(main.resolve_profile("business_management")["key"], "business_management")

    def test_computer_science_exposes_its_supported_source_layers(self) -> None:
        self.assertEqual(
            main.available_source_ids("computer_science"),
            (
                "arxiv",
                "pubmed",
                "crossref",
                "rss",
                "openalex",
                "ccf_conferences",
                "official_rss",
                "hackernews",
                "github_releases",
            ),
        )

    def test_economics_exposes_shared_user_selectable_sources(self) -> None:
        self.assertEqual(
            main.available_source_ids("economics"),
            ("arxiv", "pubmed", "openalex", "hackernews"),
        )

    def test_collect_items_keeps_legacy_empty_source_selection_collect_all_behavior(self) -> None:
        profile = {**main.resolve_profile("chemistry"), "enabled_source_ids": ()}
        since = datetime(2026, 7, 28, tzinfo=UTC)
        until = datetime(2026, 7, 29, tzinfo=UTC)
        arxiv_item = main.NewsItem("arXiv legacy result", "arXiv", None, "https://example.test/arxiv")
        pubmed_item = main.NewsItem("PubMed legacy result", "PubMed", None, "https://example.test/pubmed")
        with (
            patch.object(main, "build_session", return_value=object()),
            patch.object(main, "fetch_arxiv", return_value=[arxiv_item]),
            patch.object(main, "fetch_pubmed", return_value=[pubmed_item]),
            patch.object(main, "fetch_crossref", return_value=([], [])),
            patch.object(main, "fetch_rss", return_value=([], [])),
            patch.object(main, "fetch_openalex", return_value=[]),
        ):
            items, _ = main.collect_items(
                SimpleNamespace(source_limit=10), since, until, profile
            )

        self.assertEqual([item.title for item in items], ["arXiv legacy result", "PubMed legacy result"])

    def test_collect_items_skips_every_collector_for_an_explicit_empty_selection(self) -> None:
        profile = {
            **main.resolve_profile("chemistry"),
            "enabled_source_ids": (),
            "source_selection_explicit": True,
        }
        since = datetime(2026, 7, 28, tzinfo=UTC)
        until = datetime(2026, 7, 29, tzinfo=UTC)
        arxiv_item = main.NewsItem("must not collect", "arXiv", None, "https://example.test/arxiv")
        with (
            patch.object(main, "build_session", return_value=object()),
            patch.object(main, "fetch_arxiv", return_value=[arxiv_item]),
            patch.object(main, "fetch_pubmed", return_value=[arxiv_item]),
            patch.object(main, "fetch_crossref", return_value=([arxiv_item], [])),
            patch.object(main, "fetch_rss", return_value=([arxiv_item], [])),
            patch.object(main, "fetch_openalex", return_value=[arxiv_item]),
        ):
            items, statuses = main.collect_items(
                SimpleNamespace(source_limit=10), since, until, profile
            )

        self.assertEqual(items, [])
        self.assertEqual(statuses, [])


class FakeResponse:
    def __init__(self, payload: object) -> None:
        self.payload = payload

    def raise_for_status(self) -> None:
        return None

    def json(self) -> object:
        return self.payload


class RecordingSession:
    def __init__(self, payloads: dict[str, object]) -> None:
        self.payloads = payloads
        self.calls: list[tuple[str, dict[str, object], dict[str, object]]] = []

    def get(
        self,
        url: str,
        params: dict[str, object] | None = None,
        timeout: int = 0,
        **kwargs: object,
    ) -> FakeResponse:
        self.calls.append((url, params or {}, kwargs))
        return FakeResponse(self.payloads[url])


class PublicSourceFetcherTests(unittest.TestCase):
    def setUp(self) -> None:
        self.since = datetime(2026, 7, 28, tzinfo=UTC)
        self.until = datetime(2026, 7, 29, tzinfo=UTC)
        self.profile = main.resolve_profile("computer_science")

    def test_openalex_keeps_recent_scholarly_works_and_marks_them_academic(self) -> None:
        session = RecordingSession(
            {
                "https://api.openalex.org/works": {
                    "results": [
                        {
                            "display_name": "Data cleaning with verified repair rules",
                            "publication_date": "2026-07-28",
                            "type": "article",
                            "doi": "https://doi.org/10.1000/example",
                            "id": "https://openalex.org/W1",
                            "primary_location": {"landing_page_url": "https://doi.org/10.1000/example", "source": {"display_name": "Journal of Data Quality"}},
                            "abstract_inverted_index": {"Data": [0], "repair": [1], "study": [2]},
                            "authorships": [{"author": {"display_name": "Ada Lovelace"}}],
                        },
                        {
                            "display_name": "Old database result",
                            "publication_date": "2026-07-20",
                            "type": "article",
                            "id": "https://openalex.org/W2",
                        },
                    ]
                }
            }
        )

        items = main.fetch_openalex(session, self.since, self.until, 10, self.profile)

        self.assertEqual([item.title for item in items], ["Data cleaning with verified repair rules"])
        self.assertEqual(items[0].source_kind, "academic")
        self.assertEqual(items[0].source, "OpenAlex: Journal of Data Quality")
        self.assertEqual(items[0].authors, ["Ada Lovelace"])
        self.assertEqual(session.calls[0][1]["filter"], "from_publication_date:2026-07-28,to_publication_date:2026-07-29,type:article")

    def test_hacker_news_keeps_recent_stories_as_community_signals(self) -> None:
        session = RecordingSession(
            {
                "https://hn.algolia.com/api/v1/search_by_date": {
                    "hits": [
                        {
                            "objectID": "1",
                            "title": "Show HN: a new data repair tool",
                            "url": "https://example.test/data-repair",
                            "created_at": "2026-07-28T12:00:00.000Z",
                            "author": "researcher",
                        },
                        {
                            "objectID": "2",
                            "title": "Old agent framework",
                            "url": "https://example.test/old",
                            "created_at": "2026-07-20T12:00:00.000Z",
                        },
                    ]
                }
            }
        )

        items = main.fetch_hackernews(session, self.since, self.until, 10, self.profile)

        self.assertEqual([item.title for item in items], ["Show HN: a new data repair tool"])
        self.assertEqual(items[0].source_kind, "community")
        self.assertEqual(items[0].source, "Hacker News")
        self.assertEqual(session.calls[0][1]["tags"], "story")

    def test_github_releases_keep_recent_published_releases_as_community_signals(self) -> None:
        session = RecordingSession(
            {
                "https://api.github.com/repos/huggingface/transformers/releases": [
                    {
                        "draft": False,
                        "prerelease": False,
                        "name": "Transformers 5.0",
                        "tag_name": "v5.0.0",
                        "html_url": "https://github.com/huggingface/transformers/releases/tag/v5.0.0",
                        "body": "New agent tooling and model support.",
                        "published_at": "2026-07-28T13:00:00Z",
                        "author": {"login": "huggingface-bot"},
                    },
                    {
                        "draft": False,
                        "name": "Old release",
                        "tag_name": "v4.0.0",
                        "html_url": "https://example.test/old",
                        "published_at": "2026-07-20T13:00:00Z",
                    },
                ]
            }
        )
        profile = {**self.profile, "github_repositories": ["huggingface/transformers"]}

        items = main.fetch_github_releases(session, self.since, self.until, 10, profile)

        self.assertEqual([item.title for item in items], ["Transformers 5.0"])
        self.assertEqual(items[0].source_kind, "community")
        self.assertEqual(items[0].source, "GitHub Release: huggingface/transformers")

    def test_github_source_token_is_sent_only_to_github(self) -> None:
        session = RecordingSession({"https://api.github.com/repos/example/project/releases": []})
        profile = {"github_repositories": ["example/project"]}

        with unittest.mock.patch.dict(os.environ, {"GITHUB_SOURCE_TOKEN": "test-token"}, clear=False):
            main.fetch_github_releases(session, self.since, self.until, 10, profile)

        self.assertEqual(session.calls[0][0], "https://api.github.com/repos/example/project/releases")
        self.assertEqual(session.calls[0][2]["headers"]["Authorization"], "Bearer test-token")

    def test_europe_pmc_maps_recent_results_and_excludes_old_results(self) -> None:
        session = RecordingSession(
            {
                "https://www.ebi.ac.uk/europepmc/webservices/rest/search": {
                    "resultList": {
                        "result": [
                            {
                                "id": "PMC100",
                                "title": "Recent immune response study",
                                "journalTitle": "Open Biology",
                                "firstPublicationDate": "2026-07-28",
                                "abstractText": "A recent immune response study.",
                                "doi": "10.1000/europepmc",
                                "authorString": "Ada Lovelace; Grace Hopper",
                            },
                            {
                                "id": "PMC101",
                                "title": "Old immune response study",
                                "firstPublicationDate": "2026-07-20",
                            },
                        ]
                    }
                }
            }
        )
        profile = {**main.resolve_profile("medicine"), "relevance_terms": ["immune"]}

        items = main.fetch_europe_pmc(session, self.since, self.until, 10, profile)

        self.assertEqual([item.title for item in items], ["Recent immune response study"])
        item = items[0]
        self.assertEqual(item.source, "Europe PMC: Open Biology")
        self.assertEqual(item.link, "https://europepmc.org/article/PMC/PMC100")
        self.assertEqual(item.doi, "10.1000/europepmc")
        self.assertEqual(item.authors, ["Ada Lovelace", "Grace Hopper"])
        self.assertEqual(item.source_id, "europe_pmc")
        self.assertEqual(item.source_layer, "academic_research")
        self.assertEqual(session.calls[0][1]["format"], "json")

    def test_biorxiv_and_medrxiv_map_their_own_recent_preprints(self) -> None:
        session = RecordingSession(
            {
                "https://api.biorxiv.org/details/biorxiv/2026-07-28/2026-07-29/0": {
                    "collection": [
                        {
                            "title": "Recent bioRxiv study",
                            "date": "2026-07-28",
                            "doi": "10.1101/2026.07.28.123456",
                            "authors": "Ada Lovelace; Grace Hopper",
                            "abstract": "Recent biology preprint.",
                            "version": "1",
                        },
                        {
                            "title": "Old bioRxiv study",
                            "date": "2026-07-20",
                            "doi": "10.1101/2026.07.20.123456",
                        },
                    ]
                },
                "https://api.biorxiv.org/details/medrxiv/2026-07-28/2026-07-29/0": {
                    "collection": [
                        {
                            "title": "Recent medRxiv study",
                            "date": "2026-07-29",
                            "doi": "10.1101/2026.07.29.654321",
                            "authors": "Katherine Johnson",
                            "abstract": "Recent medical preprint.",
                            "version": "2",
                        }
                    ]
                },
            }
        )
        profile = {**main.resolve_profile("medicine"), "relevance_terms": ["biology", "medical"]}

        biorxiv_items = main.fetch_biorxiv(session, self.since, self.until, 10, profile)
        medrxiv_items = main.fetch_medrxiv(session, self.since, self.until, 10, profile)

        self.assertEqual([item.title for item in biorxiv_items], ["Recent bioRxiv study"])
        self.assertEqual(biorxiv_items[0].link, "https://www.biorxiv.org/content/10.1101/2026.07.28.123456v1")
        self.assertEqual(biorxiv_items[0].source_id, "biorxiv")
        self.assertEqual(biorxiv_items[0].source_layer, "academic_research")
        self.assertEqual([item.title for item in medrxiv_items], ["Recent medRxiv study"])
        self.assertEqual(medrxiv_items[0].link, "https://www.medrxiv.org/content/10.1101/2026.07.29.654321v2")
        self.assertEqual(medrxiv_items[0].source_id, "medrxiv")
        self.assertEqual(medrxiv_items[0].source_layer, "academic_research")

    def test_clinical_trials_maps_recent_studies_and_excludes_old_records(self) -> None:
        session = RecordingSession(
            {
                "https://clinicaltrials.gov/api/v2/studies": {
                    "studies": [
                        {
                            "protocolSection": {
                                "identificationModule": {
                                    "nctId": "NCT01234567",
                                    "briefTitle": "Recent trial of a vaccine",
                                },
                                "statusModule": {
                                    "lastUpdateSubmitDate": "2026-07-20",
                                    "lastUpdatePostDateStruct": {"date": "2026-07-28"},
                                },
                                "descriptionModule": {"briefSummary": "A randomized vaccine trial."},
                                "conditionsModule": {"conditions": ["Influenza"]},
                                "sponsorCollaboratorsModule": {
                                    "leadSponsor": {"name": "Public Health Agency"}
                                },
                            }
                        },
                        {
                            "protocolSection": {
                                "identificationModule": {"nctId": "NCT07654321", "briefTitle": "Old trial"},
                                "statusModule": {
                                    "lastUpdateSubmitDate": "2026-07-28",
                                    "lastUpdatePostDateStruct": {"date": "2026-07-20"},
                                },
                            }
                        },
                    ]
                }
            }
        )
        profile = {**main.resolve_profile("medicine"), "relevance_terms": ["vaccine", "trial"]}

        items = main.fetch_clinical_trials(session, self.since, self.until, 10, profile)

        self.assertEqual([item.title for item in items], ["Recent trial of a vaccine"])
        item = items[0]
        self.assertEqual(item.link, "https://clinicaltrials.gov/study/NCT01234567")
        self.assertEqual(item.source, "ClinicalTrials.gov")
        self.assertEqual(item.authors, ["Public Health Agency"])
        self.assertIn("Influenza", item.abstract)
        self.assertEqual(item.source_id, "clinical_trials")
        self.assertEqual(item.source_layer, "official_data_policy")
        self.assertEqual(session.calls[0][1]["pageSize"], 10)
        self.assertEqual(
            session.calls[0][1]["filter.advanced"],
            "AREA[LastUpdatePostDate]RANGE[2026-07-28,2026-07-29]",
        )
        self.assertEqual(session.calls[0][1]["sort"], "LastUpdatePostDate:desc")

    def test_academic_material_wins_over_an_equally_scored_community_signal(self) -> None:
        profile = {
            "field_keywords": {"测试": ["data"]},
            "relevance_terms": ["data"],
            "source_weights": {"OpenAlex": 50, "Hacker News": 50},
            "default_field": "测试",
        }
        academic = main.NewsItem(
            "A data result",
            "OpenAlex",
            self.until,
            "https://example.test/academic",
            abstract="data",
            source_kind="academic",
        )
        community = main.NewsItem(
            "Z data announcement",
            "Hacker News",
            self.until,
            "https://example.test/community",
            abstract="data",
            source_kind="community",
        )

        selected = main.prepare_items([community, academic], 1, self.until, profile, min_items=1)

        self.assertEqual([item.source_kind for item in selected], ["academic"])

    def test_collect_items_isolates_a_public_source_failure(self) -> None:
        profile = {
            "enabled_source_ids": ("openalex", "hackernews", "github_releases"),
            "openalex_query_terms": ["agent"],
            "community_query_terms": ["agent"],
            "github_repositories": ["example/project"],
        }
        academic = main.NewsItem("An academic agent result", "OpenAlex", self.until, "https://example.test/a")
        release = main.NewsItem(
            "Project release", "GitHub Release: example/project", self.until, "https://example.test/r", source_kind="community"
        )
        with (
            patch.object(main, "build_session", return_value=object()),
            patch.object(main, "fetch_openalex", return_value=[academic]),
            patch.object(main, "fetch_hackernews", side_effect=RuntimeError("temporary API error")),
            patch.object(main, "fetch_github_releases", return_value=[release]),
        ):
            items, statuses = main.collect_items(
                SimpleNamespace(source_limit=10), self.since, self.until, profile
            )

        self.assertEqual([item.title for item in items], ["An academic agent result", "Project release"])
        self.assertTrue(next(status for status in statuses if status.name == "Hacker News").success is False)
        self.assertEqual(next(status for status in statuses if status.name == "GitHub Releases").item_count, 1)

    def test_collect_items_isolates_new_public_source_failures_and_never_runs_unregistered_catalogue_ids(self) -> None:
        profile = {
            **main.resolve_profile("medicine"),
            "enabled_source_ids": ("europe_pmc", "biorxiv", "medrxiv", "clinical_trials", "who"),
            "source_selection_explicit": True,
        }
        europe_pmc_item = main.NewsItem(
            "Europe PMC result", "Europe PMC", self.until, "https://example.test/epmc",
            source_id="europe_pmc", source_layer="academic_research",
        )
        clinical_trials_item = main.NewsItem(
            "Trial result", "ClinicalTrials.gov", self.until, "https://example.test/trial",
            source_id="clinical_trials", source_layer="official_data_policy",
        )
        with (
            patch.object(main, "build_session", return_value=object()),
            patch.object(main, "fetch_europe_pmc", return_value=[europe_pmc_item]),
            patch.object(main, "fetch_biorxiv", side_effect=RuntimeError("temporary preprint API error")),
            patch.object(main, "fetch_medrxiv", return_value=[]),
            patch.object(main, "fetch_clinical_trials", return_value=[clinical_trials_item]),
        ):
            items, statuses = main.collect_items(
                SimpleNamespace(source_limit=10), self.since, self.until, profile
            )

        self.assertEqual([item.title for item in items], ["Europe PMC result", "Trial result"])
        self.assertFalse(next(status for status in statuses if status.source_id == "biorxiv").success)
        trial_status = next(status for status in statuses if status.source_id == "clinical_trials")
        self.assertEqual(trial_status.source_layer, "official_data_policy")
        self.assertEqual(trial_status.credibility, 5)
        self.assertNotIn("who", main.available_source_ids("medicine"))

    def test_each_new_public_collector_records_its_own_failure_without_stopping_the_task(self) -> None:
        fetchers = {
            "europe_pmc": "fetch_europe_pmc",
            "biorxiv": "fetch_biorxiv",
            "medrxiv": "fetch_medrxiv",
            "clinical_trials": "fetch_clinical_trials",
        }
        for source_id, fetcher_name in fetchers.items():
            with self.subTest(source_id=source_id):
                profile = {
                    **main.resolve_profile("medicine"),
                    "enabled_source_ids": (source_id,),
                    "source_selection_explicit": True,
                }
                with (
                    patch.object(main, "build_session", return_value=object()),
                    patch.object(main, fetcher_name, side_effect=RuntimeError("temporary public API error")),
                ):
                    items, statuses = main.collect_items(
                        SimpleNamespace(source_limit=10), self.since, self.until, profile
                    )

                self.assertEqual(items, [])
                self.assertEqual(len(statuses), 1)
                self.assertEqual(statuses[0].source_id, source_id)
                self.assertFalse(statuses[0].success)

    def test_existing_collector_items_and_statuses_have_catalogue_provenance(self) -> None:
        profile = {
            **main.resolve_profile("computer_science"),
            "enabled_source_ids": ("openalex", "hackernews", "github_releases"),
            "source_selection_explicit": True,
        }
        academic = main.NewsItem("An academic result", "OpenAlex", self.until, "https://example.test/a")
        community = main.NewsItem("A community result", "Hacker News", self.until, "https://example.test/h")
        engineering = main.NewsItem("A release", "GitHub Release", self.until, "https://example.test/g")
        with (
            patch.object(main, "build_session", return_value=object()),
            patch.object(main, "fetch_openalex", return_value=[academic]),
            patch.object(main, "fetch_hackernews", return_value=[community]),
            patch.object(main, "fetch_github_releases", return_value=[engineering]),
        ):
            items, statuses = main.collect_items(
                SimpleNamespace(source_limit=10), self.since, self.until, profile
            )

        self.assertEqual([item.source_id for item in items], ["openalex", "hackernews", "github_releases"])
        self.assertEqual(
            {status.source_id: status.source_layer for status in statuses},
            {
                "openalex": "academic_research",
                "hackernews": "community_signal",
                "github_releases": "industry_engineering",
            },
        )

    def test_document_labels_a_community_item_as_a_signal_not_academic_evidence(self) -> None:
        item = main.NewsItem(
            "Project release",
            "GitHub Release: example/project",
            self.until,
            "https://example.test/release",
            abstract="Release note.",
            source_kind="community",
            field_name="系统与软件工程",
            item_id="N001",
            comment="这是一个公开发布信号。",
        )
        profile = main.resolve_profile("computer_science")
        with TemporaryDirectory() as temporary_directory:
            output_path = main.create_document(
                [item],
                {"top_ids": ["N001"], "field_summaries": []},
                date(2026, 7, 29),
                Path(temporary_directory),
                profile,
            )
            rendered = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("社区信号", rendered)
        self.assertIn("GitHub Release: example/project", rendered)

    def test_document_does_not_place_a_community_signal_in_academic_highlights(self) -> None:
        academic = main.NewsItem(
            "Peer-reviewed data result",
            "OpenAlex: Journal of Testing",
            self.until,
            "https://example.test/paper",
            abstract="data",
            source_kind="academic",
            field_name="数据管理与时空数据",
            item_id="N001",
            comment="学术研究。",
        )
        community = main.NewsItem(
            "Community launch",
            "Hacker News",
            self.until,
            "https://example.test/community",
            abstract="data",
            source_kind="community",
            field_name="数据管理与时空数据",
            item_id="N002",
            comment="社区信号。",
        )
        with TemporaryDirectory() as temporary_directory:
            output_path = main.create_document(
                [academic, community],
                {"top_ids": ["N002"], "field_summaries": []},
                date(2026, 7, 29),
                Path(temporary_directory),
                main.resolve_profile("computer_science"),
            )
            rendered = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        highlight_text = rendered.split("分领域摘要", maxsplit=1)[0]
        self.assertIn("Peer-reviewed data result", highlight_text)
        self.assertNotIn("Community launch", highlight_text)


if __name__ == "__main__":
    unittest.main()
