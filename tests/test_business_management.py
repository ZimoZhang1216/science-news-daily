import os
import tempfile
import unittest
from datetime import date, datetime, timezone
from pathlib import Path
from unittest.mock import patch

from docx import Document

import main


class BusinessManagementProfileTests(unittest.TestCase):
    def setUp(self) -> None:
        self.profile = main.resolve_profile("business_management")

    def test_profile_is_available_with_business_output_contract(self) -> None:
        self.assertEqual(self.profile["key"], "business_management")
        self.assertEqual(self.profile["title"], "工商管理科研资讯日报")
        self.assertEqual(self.profile["output_prefix"], "business_news")
        self.assertEqual(self.profile["email_env"], "BUSINESS_REPORT_EMAIL_TO")
        self.assertFalse(self.profile["allow_default_email_fallback"])

    def test_cli_profile_help_lists_business_management(self) -> None:
        with patch.dict(os.environ, {}, clear=True):
            parser = main.build_arg_parser()
        profile_action = next(
            action for action in parser._actions if action.dest == "profile"
        )
        self.assertIn("business_management", parser.format_help())
        self.assertIn("business_management", profile_action.help)

    def test_classifies_strategy_article(self) -> None:
        field = main.classify_field(
            "Digital transformation and competitive strategy in platform firms",
            "We study how digital transformation changes competitive strategy and firm performance.",
            self.profile,
        )
        self.assertEqual(field, "战略与组织")

    def test_classifies_executive_career_research_as_hr(self) -> None:
        for title in (
            "Career mobility among executives",
            "Career advancement in organizations",
            "Executive career pathways",
        ):
            with self.subTest(title=title):
                field = main.classify_field(title, "", self.profile)
                self.assertEqual(field, "组织行为与人力资源")

        self.assertEqual(
            main.classify_field("Career transitions in the arts", "", self.profile),
            "综合工商管理",
        )

    def test_uses_only_the_36_verified_english_crossref_journals(self) -> None:
        sources = {journal["source"] for journal in self.profile["crossref_journals"]}
        self.assertEqual(len(sources), 36)
        self.assertIn("International Journal of Management Reviews", sources)
        self.assertIn("Information & Management", sources)
        self.assertNotIn("International Journal of Operations & Production Management", sources)
        self.assertNotIn("南开管理评论", sources)

    def test_rejects_non_business_management_but_keeps_management_research(self) -> None:
        excluded = main.NewsItem(
            title="Water management under climate uncertainty",
            source="arXiv",
            published=None,
            link="https://example.test/water",
            abstract="A hydrological model for reservoir operations.",
        )
        included = main.NewsItem(
            title="Algorithmic management and employee creativity in digital firms",
            source="arXiv",
            published=None,
            link="https://example.test/algorithmic",
            abstract="We examine employee behavior, leadership, and organizational innovation.",
        )
        self.assertFalse(main.is_profile_relevant(excluded, self.profile))
        self.assertTrue(main.is_profile_relevant(included, self.profile))

    def test_explicit_operations_research_can_override_a_topic_exclusion(self) -> None:
        for title in (
            "Water management in operations management",
            "Water management in service operations",
            "Water management in production operations",
        ):
            with self.subTest(title=title):
                item = main.NewsItem(
                    title=title,
                    source="Management Science",
                    published=None,
                    link="https://example.test/operations",
                )
                self.assertTrue(main.is_profile_relevant(item, self.profile))

    def test_english_term_matching_does_not_use_substrings(self) -> None:
        confirmation = main.NewsItem(
            title="Confirmation under uncertainty",
            source="Management Science",
            published=None,
            link="https://example.test/confirmation",
        )
        steam = main.NewsItem(
            title="Water management for steam systems",
            source="Management Science",
            published=None,
            link="https://example.test/steam",
        )
        journal = {"broad": False}
        self.assertFalse(
            main.should_include_crossref_item(confirmation, journal, self.profile)
        )
        self.assertFalse(main.should_include_crossref_item(steam, journal, self.profile))

    def test_safe_business_plural_forms_are_relevant_context(self) -> None:
        for title in (
            "Performance differences across firms",
            "Evidence from companies",
            "Coordination across organizations",
            "Competition across markets",
            "Decisions by managers",
            "Responses from stakeholders",
        ):
            with self.subTest(title=title):
                item = main.NewsItem(
                    title=title,
                    source="Management Science",
                    published=None,
                    link="https://example.test/plural",
                )
                self.assertTrue(main.is_profile_relevant(item, self.profile))

        for title in (
            "Water management by firms",
            "Water management by companies",
        ):
            with self.subTest(title=title):
                item = main.NewsItem(
                    title=title,
                    source="Management Science",
                    published=None,
                    link="https://example.test/plural-context",
                )
                self.assertTrue(main.is_profile_relevant(item, self.profile))

    def test_document_types_are_hard_exclusions_even_with_business_terms(self) -> None:
        for title in (
            "Editorial: innovation management",
            "Correction to corporate governance research",
            "Erratum: strategic management",
            "Book review: marketing strategy",
            "Call for papers: entrepreneurship",
            "Conference announcement: organizational behavior",
        ):
            with self.subTest(title=title):
                item = main.NewsItem(
                    title=title,
                    source="Journal of Management",
                    published=None,
                    link="https://example.test/document-type",
                )
                self.assertFalse(main.is_profile_relevant(item, self.profile))

    def test_crossref_filter_is_strict_for_business_and_unchanged_for_statistics(self) -> None:
        water_item = main.NewsItem(
            title="Water management under climate uncertainty",
            source="Management Science",
            published=None,
            link="https://example.test/water",
        )
        strategy_item = main.NewsItem(
            title="Corporate strategy and platform governance",
            source="Strategic Management Journal",
            published=None,
            link="https://example.test/strategy",
        )
        journal = {"broad": False}
        self.assertFalse(main.should_include_crossref_item(water_item, journal, self.profile))
        self.assertTrue(main.should_include_crossref_item(strategy_item, journal, self.profile))
        self.assertTrue(
            main.should_include_crossref_item(
                water_item, journal, main.resolve_profile("statistics")
            )
        )

    def test_unsupported_business_title_generalizations_fall_back(self) -> None:
        item = main.NewsItem(
            title="Competitive strategy in digital firms",
            source="Strategic Management Journal",
            published=None,
            link="https://example.test/strategy-title",
            abstract="We compare competitive strategy across digital firms.",
            field_name="战略与组织",
        )
        normalized = main.normalize_attractive_title(
            "Strategic Management Journal：材料体系揭示作用机制",
            item,
            self.profile,
        )
        self.assertNotIn("材料", normalized)
        self.assertNotIn("体系", normalized)
        self.assertNotIn("机制", normalized)
        self.assertIn("竞争战略", normalized)

    def test_materiality_and_systematic_are_not_title_evidence(self) -> None:
        item = main.NewsItem(
            title="A systematic review of ESG materiality",
            source="Management Science",
            published=None,
            link="https://example.test/materiality",
            abstract="The review compares ESG reporting priorities.",
            field_name="公司治理与可持续管理",
        )
        normalized = main.normalize_attractive_title(
            "Management Science：ESG材料体系研究",
            item,
            self.profile,
        )
        self.assertNotIn("材料", normalized)
        self.assertNotIn("体系", normalized)
        self.assertIn("ESG", normalized)

    def test_business_science_journal_prefixes_are_not_collapsed(self) -> None:
        self.assertEqual(
            main.source_title_prefix("Management Science", self.profile),
            "Management Science",
        )
        self.assertEqual(
            main.source_title_prefix("Marketing Science", self.profile),
            "Marketing Science",
        )
        self.assertEqual(main.source_title_prefix("Science", self.profile), "Science")

    def test_business_ai_title_guidance_uses_management_evidence(self) -> None:
        schema, prompt, single_item, rules = main.ai_title_guidance(self.profile)
        combined = " ".join([schema, prompt, single_item, *rules])
        self.assertIn("研究问题", combined)
        self.assertIn("理论", combined)
        self.assertIn("管理启示", combined)
        self.assertNotIn("材料/体系", combined)

    def test_business_fallback_title_uses_meaningful_terms_and_full_source(self) -> None:
        career_item = main.NewsItem(
            title="How career mobility shapes executive advancement",
            source="Human Relations",
            published=None,
            link="https://example.test/career-title",
            abstract="We study career mobility and executive career advancement.",
            field_name="组织行为与人力资源",
        )
        career_title = main.rule_based_chinese_title(career_item, self.profile)
        self.assertTrue(career_title.startswith("Human Relations："))
        self.assertNotIn("How", career_title)
        self.assertIn("职业流动", career_title)

        paradox_item = main.NewsItem(
            title="How the Paradox Shapes Entrepreneurial Ventures",
            source="Entrepreneurship Theory and Practice",
            published=None,
            link="https://example.test/paradox-title",
            abstract="The study examines an entrepreneurial paradox.",
            field_name="创新与创业",
        )
        paradox_title = main.rule_based_chinese_title(paradox_item, self.profile)
        self.assertTrue(
            paradox_title.startswith("Entrepreneurship Theory and Practice：")
        )
        self.assertNotIn("...", paradox_title)
        self.assertIn("Paradox", paradox_title)

    def test_top_item_heading_uses_actual_item_count(self) -> None:
        for count in (1, 3, 5):
            with self.subTest(count=count):
                items = [
                    main.NewsItem(
                        title=f"Business research {index}",
                        source="Journal of Management",
                        published=datetime(2026, 7, 27, tzinfo=timezone.utc),
                        link=f"https://example.test/top-{index}",
                        item_id=f"N{index:03d}",
                        field_name="战略与组织",
                        attractive_title=f"Journal of Management：竞争战略研究{index}",
                        comment="公开摘要支持对研究问题的保守整理。",
                    )
                    for index in range(1, count + 1)
                ]
                with tempfile.TemporaryDirectory() as temporary_directory:
                    output = main.create_document(
                        items,
                        {"top_ids": [item.item_id for item in items], "field_summaries": []},
                        date(2026, 7, 27),
                        Path(temporary_directory),
                        self.profile,
                    )
                    headings = [paragraph.text for paragraph in Document(output).paragraphs]
                    self.assertIn(f"今日重点 {count} 条", headings)

    def test_duplicate_top_ids_produce_one_distinct_top_item(self) -> None:
        item = main.NewsItem(
            title="Digital transformation and competitive strategy",
            source="Strategic Management Journal",
            published=datetime(2026, 7, 27, tzinfo=timezone.utc),
            link="https://example.test/duplicate-top",
            item_id="N001",
            field_name="战略与组织",
            attractive_title="Strategic Management Journal：数字化转型与竞争战略",
            comment="公开摘要支持对研究问题的保守整理。",
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = main.create_document(
                [item],
                {"top_ids": ["N001"] * 5, "field_summaries": []},
                date(2026, 7, 27),
                Path(temporary_directory),
                self.profile,
            )
            paragraphs = [paragraph.text for paragraph in Document(output).paragraphs]
            self.assertIn("今日重点 1 条", paragraphs)
            self.assertEqual(
                sum(text.startswith("01  ") for text in paragraphs),
                1,
            )
            self.assertFalse(any(text.startswith("02  ") for text in paragraphs))

    def test_generated_document_uses_platform_cjk_font_for_runs_and_styles(self) -> None:
        item = main.NewsItem(
            title="Digital transformation and competitive strategy",
            source="Strategic Management Journal",
            published=datetime(2026, 7, 27, tzinfo=timezone.utc),
            link="https://example.test/cjk-font",
            item_id="N001",
            field_name="战略与组织",
            attractive_title="Strategic Management Journal：数字化转型与竞争战略",
            comment="公开摘要支持对研究问题的保守整理。",
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = main.create_document(
                [item],
                {"top_ids": ["N001"], "field_summaries": []},
                date(2026, 7, 27),
                Path(temporary_directory),
                self.profile,
            )
            document = Document(output)

        expected = main.resolve_cjk_font()
        for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
            fonts = document.styles[style_name]._element.rPr.rFonts
            self.assertEqual(fonts.get(main.qn("w:eastAsia")), expected)
            self.assertEqual(fonts.get(main.qn("w:ascii")), "Arial")
            self.assertEqual(fonts.get(main.qn("w:hAnsi")), "Arial")

        cjk_run_fonts = [
            run._element.rPr.rFonts.get(main.qn("w:eastAsia"))
            for paragraph in document.paragraphs
            for run in paragraph.runs
            if any("\u4e00" <= character <= "\u9fff" for character in run.text)
            and run._element.rPr is not None
            and run._element.rPr.rFonts is not None
        ]
        self.assertIn(expected, cjk_run_fonts)

    def test_cjk_font_selection_covers_declared_platforms(self) -> None:
        self.assertEqual(main.resolve_cjk_font("darwin"), "Hiragino Sans GB")
        self.assertEqual(main.resolve_cjk_font("linux"), "Noto Sans CJK SC")
        self.assertEqual(main.resolve_cjk_font("win32"), "Microsoft YaHei")

    def test_document_uses_the_business_output_prefix(self) -> None:
        item = main.NewsItem(
            title="Digital transformation and competitive strategy",
            source="Strategic Management Journal",
            published=datetime(2026, 7, 27, tzinfo=timezone.utc),
            link="https://doi.org/10.1000/example",
            abstract="A public abstract about digital transformation and competitive strategy.",
            item_id="N001",
            field_name="战略与组织",
            attractive_title="Strategic Management Journal：数字化转型与竞争战略",
            comment="当前可获取信息仅支持对研究对象与公开摘要的保守整理。",
        )
        with tempfile.TemporaryDirectory() as temporary_directory:
            output = main.create_document(
                [item],
                {"top_ids": ["N001"], "field_summaries": []},
                date(2026, 7, 27),
                Path(temporary_directory),
                self.profile,
            )
            self.assertEqual(output.name, "business_news_2026-07-27.docx")
            self.assertTrue(output.exists())
            self.assertGreater(output.stat().st_size, 0)


if __name__ == "__main__":
    unittest.main()
